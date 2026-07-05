"""PdfToWordMixin — PDF → Word/DOCX conversion methods."""

import os
import tempfile
import shutil
from pathlib import Path
from datetime import datetime

from PySide6.QtWidgets import (QWidget, QVBoxLayout, QLabel, QMessageBox, QGroupBox,
                               QLineEdit, QDialog, QDialogButtonBox, QTabWidget, QScrollArea)
from PySide6.QtCore import Qt

try:
    from pdf2docx import Converter as _Converter
    Converter = _Converter
except ImportError as _e:
    Converter = None
    print(f"[IMPORT] pdf2docx not available: {_e}")

try:
    from pypdf import PdfReader
except ImportError as _e:
    PdfReader = None
    print(f"[IMPORT] pypdf not available: {_e}")

try:
    from docx import Document
    from docx.shared import Inches
except ImportError as _e:
    Document = Inches = None
    print(f"[IMPORT] python-docx not available: {_e}")

from widgets import AnimatedCheckBox
from dialogs import PdfToWordDialog


class PdfToWordMixin:
    """Mixin: PDF → Word/DOCX conversion for FileConverterApp."""

    def check_pdf_has_images(self, pdf_path):
        try:
            import fitz
            doc = fitz.open(pdf_path)
            has_images = False
            for page in doc:
                if page.get_images():
                    has_images = True
                    break
            doc.close()
            return has_images
        except Exception:
            return False

    def handle_encrypted_pdfs(self, encrypted_pdfs, all_pdfs):
        dialog = QDialog(self)
        dialog.setWindowTitle(self.translate_text("PDFs protégés par mot de passe"))
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout(dialog)

        if len(encrypted_pdfs) == 1:
            message = self.translate_text(f"Le fichier '{Path(encrypted_pdfs[0]).name}' est protégé par mot de passe.")
        else:
            message = self.translate_text(f"{len(encrypted_pdfs)} fichiers PDF sont protégés par mot de passe.")

        message += "\n\n" + self.translate_text("Pour convertir ces fichiers en Word, vous devez fournir les mots de passe.\n\n")
        message += self.translate_text("Options :")

        info_label = QLabel(message)
        info_label.setWordWrap(True)
        info_label.setStyleSheet("font-weight: bold; color: #d35400; padding: 10px;")
        layout.addWidget(info_label)

        tab_widget = QTabWidget()

        single_tab = QWidget()
        single_layout = QVBoxLayout(single_tab)

        if len(encrypted_pdfs) == 1:
            single_layout.addWidget(QLabel(self.translate_text(f"Fichier : {Path(encrypted_pdfs[0]).name}")))
            single_password_input = QLineEdit()
            single_password_input.setEchoMode(QLineEdit.Password)
            single_password_input.setPlaceholderText(self.translate_text("Mot de passe du PDF"))

            single_layout.addWidget(QLabel(self.translate_text("Mot de passe :")))
            single_layout.addWidget(single_password_input)

            dialog.single_password_input = single_password_input
            dialog.single_pdf = encrypted_pdfs[0]

        multiple_tab = QWidget()
        multiple_layout = QVBoxLayout(multiple_tab)

        if len(encrypted_pdfs) > 1:
            scroll_area = QScrollArea()
            scroll_widget = QWidget()
            scroll_layout = QVBoxLayout(scroll_widget)

            password_inputs = {}

            for pdf_file in encrypted_pdfs:
                group = QGroupBox(Path(pdf_file).name)
                group_layout = QVBoxLayout(group)

                password_input = QLineEdit()
                password_input.setEchoMode(QLineEdit.Password)
                password_input.setPlaceholderText(self.translate_text("Mot de passe (laisser vide si aucun)"))

                group_layout.addWidget(QLabel(self.translate_text("Mot de passe :")))
                group_layout.addWidget(password_input)

                password_inputs[pdf_file] = password_input
                scroll_layout.addWidget(group)

            scroll_widget.setLayout(scroll_layout)
            scroll_area.setWidget(scroll_widget)
            scroll_area.setWidgetResizable(True)
            multiple_layout.addWidget(scroll_area)

            dialog.password_inputs = password_inputs

        same_tab = QWidget()
        same_layout = QVBoxLayout(same_tab)

        same_password_input = QLineEdit()
        same_password_input.setEchoMode(QLineEdit.Password)
        same_password_input.setPlaceholderText(self.translate_text("Mot de passe commun à tous les PDF"))

        same_layout.addWidget(QLabel(self.translate_text("Ce mot de passe sera essayé sur tous les PDF chiffrés.")))
        same_layout.addWidget(QLabel(self.translate_text("Mot de passe commun :")))
        same_layout.addWidget(same_password_input)

        dialog.same_password_input = same_password_input

        if len(encrypted_pdfs) == 1:
            tab_widget.addTab(single_tab, self.translate_text("PDF unique"))
        else:
            tab_widget.addTab(multiple_tab, self.translate_text("PDF multiples"))
            tab_widget.addTab(same_tab, self.translate_text("Mot de passe commun"))

        layout.addWidget(tab_widget)

        options_group = QGroupBox(self.translate_text("Options de sortie"))
        options_layout = QVBoxLayout(options_group)

        keep_encrypted_check = AnimatedCheckBox(
            self.translate_text("Conserver les versions originales chiffrées après conversion")
        )
        keep_encrypted_check.setChecked(True)

        remove_passwords_check = AnimatedCheckBox(
            self.translate_text("Supprimer la protection par mot de passe des PDF convertis (recommandé)")
        )
        remove_passwords_check.setChecked(True)

        options_layout.addWidget(keep_encrypted_check)
        options_layout.addWidget(remove_passwords_check)

        layout.addWidget(options_group)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)

        def collect_passwords():
            passwords = {}

            if len(encrypted_pdfs) == 1:
                password = dialog.single_password_input.text()
                if password:
                    passwords[dialog.single_pdf] = password
            else:
                if tab_widget.currentIndex() == 0:
                    for pdf_file, password_input in dialog.password_inputs.items():
                        password = password_input.text()
                        if password:
                            passwords[pdf_file] = password
                else:
                    common_password = dialog.same_password_input.text()
                    if common_password:
                        for pdf_file in encrypted_pdfs:
                            passwords[pdf_file] = common_password

            dialog.passwords = passwords
            dialog.keep_originals = keep_encrypted_check.isChecked()
            dialog.remove_passwords = remove_passwords_check.isChecked()
            dialog.accept()

        button_box.accepted.connect(collect_passwords)
        button_box.rejected.connect(dialog.reject)

        layout.addWidget(button_box)

        if dialog.exec() == QDialog.Accepted:
            if hasattr(dialog, 'passwords') and dialog.passwords:
                return {
                    'passwords': dialog.passwords,
                    'keep_originals': dialog.keep_originals,
                    'remove_passwords': dialog.remove_passwords
                }
            else:
                QMessageBox.warning(self, self.translate_text("Avertissement"),
                                self.translate_text("Aucun mot de passe fourni. La conversion sera annulée."))
                return False
        else:
            return False

    def decrypt_pdfs_before_conversion(self, pdf_files, passwords_info):
        temp_files = []
        success_count = 0
        failed_files = []

        passwords = passwords_info.get('passwords', {})
        keep_originals = passwords_info.get('keep_originals', True)
        remove_passwords = passwords_info.get('remove_passwords', True)

        self.show_progress(True, self.translate_text("Déchiffrement des PDF..."))

        for i, pdf_file in enumerate(pdf_files):
            try:
                with open(pdf_file, 'rb') as f:
                    pdf_reader = PdfReader(f)

                    if pdf_reader.is_encrypted:
                        password = passwords.get(pdf_file)

                        if not password:
                            try:
                                pdf_reader.decrypt('')
                            except Exception:
                                failed_files.append(f"{Path(pdf_file).name} - {self.translate_text('Mot de passe requis')}")
                                continue
                        else:
                            success = pdf_reader.decrypt(password)

                            if not success:
                                for pwd_variation in [password, password.lower(), password.upper()]:
                                    try:
                                        success = pdf_reader.decrypt(pwd_variation)
                                        if success:
                                            break
                                    except Exception:
                                        continue

                                if not success:
                                    failed_files.append(f"{Path(pdf_file).name} - {self.translate_text('Mot de passe incorrect')}")
                                    continue

                        pdf_writer = __import__('pypdf', fromlist=['PdfWriter']).PdfWriter()

                        for page in pdf_reader.pages:
                            pdf_writer.add_page(page)

                        original_stem = Path(pdf_file).stem
                        temp_file = os.path.join(
                            tempfile.gettempdir(),
                            f"{original_stem}_decrypted.pdf"
                        )

                        counter = 1
                        while os.path.exists(temp_file):
                            temp_file = os.path.join(
                                tempfile.gettempdir(),
                                f"{original_stem}_decrypted_{counter}.pdf"
                            )
                            counter += 1
                        self.temp_files.append(temp_file)

                        with open(temp_file, 'wb') as output_file:
                            if remove_passwords:
                                pdf_writer.write(output_file)
                            else:
                                pdf_writer.encrypt(password)
                                pdf_writer.write(output_file)

                        temp_files.append(temp_file)
                        success_count += 1

                        if not keep_originals:
                            backup_file = pdf_file + ".backup"
                            shutil.copy2(pdf_file, backup_file)
                            shutil.copy2(temp_file, pdf_file)
                            temp_files.append(backup_file)
                    else:
                        temp_files.append(pdf_file)
                        success_count += 1

            except Exception as e:
                error_msg = f"{Path(pdf_file).name} - {str(e)}"
                failed_files.append(error_msg)
                print(f"Decryption error {pdf_file}: {e}")

            self.progress_bar.setValue(int((i + 1) / len(pdf_files) * 100))

        self.show_progress(False)

        if failed_files:
            error_message = self.translate_text(f"{success_count}/{len(pdf_files)} PDF(s) déchiffré(s) avec succès.\n\n")
            error_message += self.translate_text(f"Échecs ({len(failed_files)}):\n")
            error_message += "\n".join(failed_files[:3])

            if len(failed_files) > 3:
                error_message += f"\n... et {len(failed_files) - 3} autres"

            QMessageBox.warning(self, self.translate_text("Résultat déchiffrement"), error_message)

        return temp_files

    def decrypt_single_pdf(self, pdf_path, password, output_path=None, remove_password=True):
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PdfReader(f)

                if not pdf_reader.is_encrypted:
                    return pdf_path

                success = pdf_reader.decrypt(password)

                if not success:
                    for pwd_variation in [password, password.lower(), password.upper()]:
                        try:
                            success = pdf_reader.decrypt(pwd_variation)
                            if success:
                                break
                        except Exception:
                            continue

                    if not success:
                        raise Exception(self.translate_text("Mot de passe incorrect"))

                pdf_writer = __import__('pypdf', fromlist=['PdfWriter']).PdfWriter()

                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)

                if not output_path:
                    output_path = self.create_temp_file(suffix="_decrypted.pdf")

                with open(output_path, 'wb') as output_file:
                    if remove_password:
                        pdf_writer.write(output_file)
                    else:
                        pdf_writer.encrypt(password)
                        pdf_writer.write(output_file)

                return output_path

        except Exception as e:
            print(f"Decryption error {pdf_path}: {e}")
            raise

    def convert_pdfs_without_images(self, pdf_files):
        output_dir = self.get_output_directory()
        if not output_dir:
            return

        message = self.translate_text("conversion_pdf_to_word").format(len(pdf_files))
        self.show_progress(True, message)

        success_count = 0
        start_time = datetime.now()

        for i, file_path in enumerate(pdf_files):
            try:
                output_file = os.path.join(output_dir, f"{Path(file_path).stem}.docx")
                file_size = os.path.getsize(file_path)

                operation_start = datetime.now()

                current_mode = self.config.get("pdf_to_word_mode", "with_images")
                if current_mode == "with_images":
                    self.convert_pdf_to_docx_improved(file_path, output_file)
                elif current_mode == "text_only":
                    self.convert_pdf_to_docx_text_only(file_path, output_file)
                else:
                    self.convert_pdf_to_docx_text_only(file_path, output_file)

                operation_time = (datetime.now() - operation_start).total_seconds()
                self.achievement_system.record_conversion("pdf_to_word", file_size, True)
                self.achievement_system.mark_format_as_used("pdf")
                self.achievement_system.mark_format_as_used("docx")

                self.db_manager.add_conversion_record(
                    source_file=file_path,
                    source_format="PDF",
                    target_file=output_file,
                    target_format="DOCX",
                    operation_type="pdf_to_word",
                    file_size=file_size,
                    conversion_time=operation_time,
                    success=True,
                    notes=f"Mode: {current_mode}"
                )

                success_count += 1
                self.progress_bar.setValue(int((i + 1) / len(pdf_files) * 100))

            except Exception as e:
                self.db_manager.add_conversion_record(
                    source_file=file_path,
                    source_format="PDF",
                    target_file="",
                    target_format="DOCX",
                    operation_type="pdf_to_word",
                    file_size=0,
                    conversion_time=0,
                    success=False,
                    notes=f"Error: {str(e)}"
                )
                QMessageBox.critical(self, self.translate_text("Erreur"), self.translate_text(f"Erreur avec {Path(file_path).name}: {str(e)}"))

        total_time = (datetime.now() - start_time).total_seconds()
        self.show_progress(False)

        if success_count > 0 and 0 <= datetime.now().hour < 6:
            self.achievement_system.increment_stat("night_conversions", success_count)
            self.achievement_system.check_achievement("night_owl")

        if success_count >= 50 and total_time <= 300:
            self.achievement_system.update_stat("recent_batch_files", success_count)
            self.achievement_system.update_stat("recent_batch_time", total_time)
            self.achievement_system.check_speed_conversion(success_count, total_time)

        formatted_time = f"{total_time:.1f}"
        message = self.translate_text("pdf_to_word_success_sum").format(success_count,len(pdf_files),formatted_time)
        if self.config.get("enable_system_notifications", True):
            self.system_notifier.send("pdf_to_word")
        QMessageBox.information(self, self.translate_text("Succès"), self.translate_text(message))

    def _get_chart_regions(self, page):
        MIN_PATHS   = 8
        MIN_AREA    = 8000
        PADDING     = 10
        MERGE_GAP   = 20

        import fitz
        drawings = page.get_drawings()
        if not drawings:
            return []

        path_rects = [fitz.Rect(d["rect"]) for d in drawings if d.get("rect")]
        if len(path_rects) < MIN_PATHS:
            return []

        path_rects.sort(key=lambda r: (r.y0, r.x0))
        clusters = []
        cur = fitz.Rect(path_rects[0])
        count = 1

        for r in path_rects[1:]:
            expanded = fitz.Rect(cur.x0 - MERGE_GAP, cur.y0 - MERGE_GAP,
                                 cur.x1 + MERGE_GAP, cur.y1 + MERGE_GAP)
            if expanded.intersects(r):
                cur = fitz.Rect(
                    min(cur.x0, r.x0), min(cur.y0, r.y0),
                    max(cur.x1, r.x1), max(cur.y1, r.y1)
                )
                count += 1
            else:
                if count >= MIN_PATHS and cur.get_area() >= MIN_AREA:
                    clusters.append(cur)
                cur = fitz.Rect(r)
                count = 1

        if count >= MIN_PATHS and cur.get_area() >= MIN_AREA:
            clusters.append(cur)

        page_rect = page.rect
        result = []
        for c in clusters:
            padded = fitz.Rect(
                max(0,            c.x0 - PADDING),
                max(0,            c.y0 - PADDING),
                min(page_rect.x1, c.x1 + PADDING),
                min(page_rect.y1, c.y1 + PADDING),
            )
            result.append(padded)

        return result

    def _rasterize_region(self, page, rect, dpi=150):
        import fitz
        zoom   = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        fitz.IRect(
            int(rect.x0 * zoom), int(rect.y0 * zoom),
            int(rect.x1 * zoom), int(rect.y1 * zoom),
        )
        pix = page.get_pixmap(matrix=matrix, clip=rect, alpha=False)
        return pix.tobytes("png")

    def _try_word_com_pdf_to_docx(self, pdf_path, docx_path):
        import sys
        if sys.platform != "win32":
            return False

        try:
            import winreg
            winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, "Word.Application")
        except (OSError, ImportError):
            print("[PDF→DOCX] Word not found in registry — skipping COM tier")
            return False

        import os
        import threading
        pdf_path  = os.path.abspath(pdf_path)
        docx_path = os.path.abspath(docx_path)

        result      = {"ok": False}
        stop_event  = threading.Event()

        def _dialog_dismisser():
            import ctypes
            import ctypes.wintypes
            user32   = ctypes.windll.user32
            BM_CLICK = 0x00F5

            DialogProc = ctypes.WINFUNCTYPE(
                ctypes.c_bool,
                ctypes.wintypes.HWND,
                ctypes.wintypes.LPARAM,
            )

            def _find_ok_button(hwnd_dialog):
                found = ctypes.wintypes.HWND(0)

                @DialogProc
                def _enum(hwnd_child, _lp):
                    buf = ctypes.create_unicode_buffer(64)
                    user32.GetWindowTextW(hwnd_child, buf, 64)
                    if buf.value.strip().upper() in ("OK", "O&K"):
                        found.value = hwnd_child
                        return False
                    ctrl_id = user32.GetDlgCtrlID(hwnd_child)
                    if ctrl_id == 1:
                        found.value = hwnd_child
                        return False
                    return True

                user32.EnumChildWindows(hwnd_dialog, _enum, 0)
                return found.value or None

            dialog_classes = ["#32770", "bosa_sdm_msword"]
            target_title   = "Microsoft Word"

            while not stop_event.is_set():
                for cls in dialog_classes:
                    hwnd = user32.FindWindowW(cls, None)
                    if hwnd:
                        ok_btn = _find_ok_button(hwnd)
                        if ok_btn:
                            user32.SendMessageW(ok_btn, BM_CLICK, 0, 0)
                        else:
                            user32.SendMessageW(hwnd, 0x0111, 1, 0)
                hwnd = user32.FindWindowW(None, target_title)
                if hwnd:
                    ok_btn = _find_ok_button(hwnd)
                    if ok_btn:
                        user32.SendMessageW(ok_btn, BM_CLICK, 0, 0)
                    else:
                        user32.SendMessageW(hwnd, 0x0111, 1, 0)
                import time as _t
                _t.sleep(0.3)

        def _worker():
            word = None
            doc  = None
            try:
                import pythoncom
                import comtypes.client
                pythoncom.CoInitialize()
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible        = False
                word.DisplayAlerts  = 0
                word.AutomationSecurity = 3

                doc = word.Documents.Open(
                    pdf_path,
                    ConfirmConversions = False,
                    ReadOnly           = True,
                    AddToRecentFiles   = False,
                    NoEncodingDialog   = True,
                )
                doc.SaveAs2(docx_path, FileFormat=16)
                result["ok"] = True
                print(f"[PDF→DOCX] Word COM success: {os.path.basename(pdf_path)}")
            except Exception as e:
                print(f"[PDF→DOCX] Word COM failed: {e}")
            finally:
                if doc is not None:
                    try: doc.Close(False)
                    except Exception: pass
                if word is not None:
                    try: word.Quit()
                    except Exception: pass
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except Exception: pass

        dismisser = threading.Thread(target=_dialog_dismisser, daemon=True)
        dismisser.start()

        t = threading.Thread(target=_worker, daemon=True)
        t.start()
        t.join(timeout=60)

        stop_event.set()

        if t.is_alive():
            print("[PDF→DOCX] Word COM timed out — falling back to pdf2docx")
            return False

        return result["ok"]

    def _try_pdf2docx(self, pdf_path, docx_path):
        if Converter is None:
            print("[PDF→DOCX] pdf2docx not available — skipping Tier 2")
            return False
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None, parse_drawing=True)
            cv.close()
            print(f"[PDF→DOCX] pdf2docx success: {Path(pdf_path).name}")
            return True
        except Exception as e:
            print(f"[PDF→DOCX] pdf2docx failed: {e}")
            return False

    def _solid_fitz_fallback(self, pdf_path, docx_path):
        import io as _io
        import fitz
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        from docx.shared import Cm
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        pdf_doc = fitz.open(pdf_path)

        def _flag_to_style(flags, size, median_size):
            is_bold   = bool(flags & 16)
            is_italic = bool(flags & 2)
            ratio     = size / median_size if median_size else 1.0
            if ratio >= 1.6 and is_bold:
                return "heading1"
            if ratio >= 1.3 and is_bold:
                return "heading2"
            if ratio >= 1.1 and is_bold:
                return "heading3"
            return ("bold_body" if is_bold else
                    "italic_body" if is_italic else "body")

        def _apply_run_style(run, flags, size):
            run.bold   = bool(flags & 16)
            run.italic = bool(flags & 2)
            if size:
                run.font.size = Pt(round(size, 1))

        def _add_table_to_doc(doc, tab):
            try:
                rows_data = tab.extract()
                if not rows_data or not rows_data[0]:
                    return
                n_rows = len(rows_data)
                n_cols = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=n_rows, cols=n_cols)
                tbl.style = "Table Grid"
                for r_idx, row_data in enumerate(rows_data):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < n_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = (cell_text or "").strip()
                doc.add_paragraph()
            except Exception as te:
                print(f"[fallback table] {te}")

        for page_num in range(len(pdf_doc)):
            if page_num > 0:
                doc.add_page_break()

            page     = pdf_doc.load_page(page_num)
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            all_sizes = []
            for blk in text_dict.get("blocks", []):
                if blk["type"] != 0:
                    continue
                for line in blk.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("size", 0) > 0:
                            all_sizes.append(span["size"])
            all_sizes.sort()
            median_size = all_sizes[len(all_sizes) // 2] if all_sizes else 11.0

            table_rects = []
            tables_on_page = []
            try:
                finder = page.find_tables()
                tables_on_page = finder.tables if finder else []
                for t in tables_on_page:
                    table_rects.append(fitz.Rect(t.bbox))
            except Exception:
                pass

            def _rect_in_table(rect):
                for tr in table_rects:
                    inter = fitz.Rect(rect).intersect(tr)
                    if not inter.is_empty and inter.get_area() > 100:
                        return True
                return False

            images_on_page = []
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    bbox_list = page.get_image_rects(xref)
                    if not bbox_list:
                        continue
                    bbox = bbox_list[0]
                    if _rect_in_table(bbox):
                        continue
                    pix = fitz.Pixmap(pdf_doc, xref)
                    if pix.colorspace and pix.colorspace.n > 3:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    if pix.width < 10 or pix.height < 10:
                        continue
                    img_bytes = pix.tobytes("png")
                    images_on_page.append({
                        "y0": bbox.y0,
                        "y1": bbox.y1,
                        "bytes": img_bytes,
                        "width_pt": bbox.width,
                    })
                    pix = None
                except Exception as ie:
                    print(f"[fallback img xref={xref}] {ie}")

            images_on_page.sort(key=lambda x: x["y0"])

            stream = []
            for blk in text_dict.get("blocks", []):
                y0 = blk.get("bbox", [0, 0])[1]
                if blk["type"] == 0:
                    if not _rect_in_table(fitz.Rect(blk["bbox"])):
                        stream.append(("text", y0, blk))
            for img in images_on_page:
                stream.append(("image", img["y0"], img))
            for tab in tables_on_page:
                try:
                    stream.append(("table", tab.bbox[1], tab))
                except Exception:
                    pass

            stream.sort(key=lambda x: x[1])

            img_inserted_xranges = set()

            for item_type, y0, obj in stream:

                if item_type == "image":
                    key = round(y0 / 5) * 5
                    if key in img_inserted_xranges:
                        continue
                    img_inserted_xranges.add(key)
                    try:
                        buf = _io.BytesIO(obj["bytes"])
                        w_pt  = obj["width_pt"]
                        max_w = Inches(5.5)
                        width = min(Pt(w_pt), max_w)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(buf, width=width)
                    except Exception as ie:
                        print(f"[fallback img render] {ie}")

                elif item_type == "table":
                    _add_table_to_doc(doc, obj)

                elif item_type == "text":
                    blk = obj
                    for line in blk.get("lines", []):
                        spans = line.get("spans", [])
                        if not spans:
                            continue

                        line_text = "".join(s.get("text", "") for s in spans)
                        if not line_text.strip():
                            continue
                        dom = max(spans, key=lambda s: len(s.get("text", "")))
                        flags = dom.get("flags", 0)
                        size  = dom.get("size", median_size)
                        style_label = _flag_to_style(flags, size, median_size)

                        if style_label == "heading1":
                            p = doc.add_heading(line_text.strip(), level=1)
                        elif style_label == "heading2":
                            p = doc.add_heading(line_text.strip(), level=2)
                        elif style_label == "heading3":
                            p = doc.add_heading(line_text.strip(), level=3)
                        else:
                            p = doc.add_paragraph()
                            for span in spans:
                                span_text = span.get("text", "")
                                if not span_text:
                                    continue
                                run = p.add_run(span_text)
                                _apply_run_style(run, span.get("flags", 0),
                                                 span.get("size", None))

        pdf_doc.close()
        doc.save(docx_path)
        print(f"[PDF→DOCX] solid fitz fallback done: {Path(docx_path).name}")

    def convert_pdf_to_docx_with_charts(self, pdf_path, docx_path):
        import io as _io
        import fitz

        if self._try_word_com_pdf_to_docx(pdf_path, docx_path):
            return

        base_ok = self._try_pdf2docx(pdf_path, docx_path)

        if base_ok:
            try:
                pdf_doc = fitz.open(pdf_path)
                docx    = Document(docx_path)

                charts_per_page = {}
                for page_num in range(len(pdf_doc)):
                    page  = pdf_doc.load_page(page_num)
                    rects = self._get_chart_regions(page)
                    if not rects:
                        continue
                    imgs = []
                    for rect in rects:
                        png_bytes = self._rasterize_region(page, rect, dpi=150)
                        imgs.append(_io.BytesIO(png_bytes))
                    charts_per_page[page_num] = imgs

                pdf_doc.close()

                if charts_per_page:
                    from docx.shared import Inches as _Inches
                    import copy

                    def _add_picture_paragraph(docx_doc, img_buf, max_width_inches=5.5):
                        p   = docx_doc.add_paragraph()
                        p.alignment = 1
                        run = p.add_run()
                        run.add_picture(img_buf, width=_Inches(max_width_inches))
                        return p

                    def _insert_para_after(ref_para, new_para_xml):
                        ref_para._element.addnext(new_para_xml)

                    body_paras = docx.paragraphs
                    page_break_indices = []
                    for idx, para in enumerate(body_paras):
                        for run in para.runs:
                            if (run._element.xml.find("w:lastRenderedPageBreak") != -1 or
                                    run._element.xml.find("w:br") != -1):
                                page_break_indices.append(idx)
                                break

                    for page_num in sorted(charts_per_page.keys()):
                        imgs = charts_per_page[page_num]
                        if page_num == 0 or page_num - 1 >= len(page_break_indices):
                            for img_buf in imgs:
                                _add_picture_paragraph(docx, img_buf)
                        else:
                            ref_idx  = page_break_indices[page_num - 1]
                            ref_para = body_paras[ref_idx]
                            for img_buf in reversed(imgs):
                                tmp_doc  = Document()
                                pic_para = _add_picture_paragraph(tmp_doc, img_buf)
                                _insert_para_after(ref_para, copy.deepcopy(pic_para._element))

                    docx.save(docx_path)
            except Exception as e:
                print(f"[chart injection] {e}")
            return

        print(f"[PDF→DOCX] falling back to solid fitz for {Path(pdf_path).name}")
        self._solid_fitz_fallback(pdf_path, docx_path)

    def convert_pdf_to_docx_improved(self, pdf_path, docx_path):
        self.convert_pdf_to_docx_with_charts(pdf_path, docx_path)

    def convert_pdf_to_docx_text_only(self, pdf_path, docx_path):
        try:
            import fitz
            doc = Document()
            pdf_document = fitz.open(pdf_path)
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                if text.strip():
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            pdf_document.close()
            doc.save(docx_path)
        except Exception:
            self._solid_fitz_fallback(pdf_path, docx_path)

    def convert_pdf_to_docx_fallback(self, pdf_path, docx_path):
        self._solid_fitz_fallback(pdf_path, docx_path)

    def convert_pdf_to_docx_basic(self, pdf_path, docx_path):
        try:
            self._solid_fitz_fallback(pdf_path, docx_path)
        except Exception:
            import fitz
            doc = Document()
            pdf_document = fitz.open(pdf_path)
            doc.add_heading(f"Conversion de: {Path(pdf_path).name}", level=1)
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                doc.add_paragraph(f"--- Page {page_num + 1} ---")
                doc.add_paragraph(text if text.strip() else "[Graphical content]")
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            pdf_document.close()
            doc.save(docx_path)

    def convert_pdf_to_word(self):
        if not (hasattr(self, 'active_templates') and 'pdf_to_word' in self.active_templates):
            _def_id, _ = (self._ensure_template_manager() or object()).get_default_template('Conversion PDF→Word')
            if _def_id:
                (self._ensure_template_manager() or object()).apply_template(_def_id, self)

        if hasattr(self, 'active_templates') and 'word_to_pdf' in self.active_templates:
            template = self.active_templates['word_to_pdf']
            template.get('page_format_value', 'a4')
            template.get('orientation_value', 'portrait')
            template.get('quality_value', 150)
            template.get('include_metadata', True)
            template.get('compress', True)

        selected_items = self.files_list_widget.selectedItems()
        files_to_process = []
        if selected_items:
            for i in range(self.files_list_widget.count()):
                item = self.files_list_widget.item(i)
                if item.isSelected():
                    files_to_process.append(item.data(Qt.UserRole))
        else:
            files_to_process = [f for f in self.files_list if f.lower().endswith('.pdf')]

        pdf_files = [f for f in files_to_process if f.lower().endswith('.pdf')]
        if not pdf_files:
            msg = self.translate_text("Veuillez sélectionner au moins un fichier PDF") if selected_items else self.translate_text("La liste doit contenir au moins un fichier PDF")
            QMessageBox.warning(self, self.translate_text("Avertissement"), msg)
            return

        encrypted_pdfs = []
        for pdf_file in pdf_files:
            try:
                with open(pdf_file, 'rb') as f:
                    pdf_reader = PdfReader(f)
                    if pdf_reader.is_encrypted:
                        encrypted_pdfs.append(pdf_file)
            except Exception as e:
                print(f"PDF verification error {pdf_file}: {e}")

        if encrypted_pdfs:
            response = self.handle_encrypted_pdfs(encrypted_pdfs, pdf_files)
            if response is False:
                return
            elif response is not None:
                passwords = response
                pdf_files = self.decrypt_pdfs_before_conversion(pdf_files, passwords)
                print(f"[DEBUG] {len(encrypted_pdfs)} password-protected PDF files successfully converted.")
                self.achievement_system.record_protected_file_conversion(len(encrypted_pdfs), "pdf")
                if not pdf_files:
                    return

        pdfs_with_images = []
        for pdf_file in pdf_files:
            if self.check_pdf_has_images(pdf_file):
                pdfs_with_images.append(pdf_file)

        if not pdfs_with_images:
            self.convert_pdfs_without_images(pdf_files)
            return

        if hasattr(self, 'active_templates') and 'pdf_to_word' in self.active_templates:
            conversion_mode = self.active_templates['pdf_to_word'].get('mode', 'with_images')
        else:
            current_mode = self.config.get("pdf_to_word_mode", "with_images")
            dialog = PdfToWordDialog(self, self.current_language, current_mode, has_images=True)
            if dialog.exec() != QDialog.Accepted:
                return
            conversion_mode = dialog.get_conversion_mode()

        output_dir = self.get_output_directory()
        if not output_dir:
            return

        message = self.translate_text("conversion_pdf_to_word").format(len(pdf_files))
        self.show_progress(True, message)
        self._set_ui_enabled(False)

        _mode   = conversion_mode
        _outdir = output_dir

        def _run_pdf_to_word(task):
            import os, time as _time
            t0 = _time.perf_counter()
            fp = task["input_path"]
            out = task["output_path"]
            fs  = os.path.getsize(fp) if os.path.exists(fp) else 0
            if _mode == "with_images":
                self.convert_pdf_to_docx_improved(fp, out)
            elif _mode == "text_only":
                self.convert_pdf_to_docx_text_only(fp, out)
            else:
                self.convert_pdf_to_docx_with_image_text(fp, out)
            return {"success": True, "error": "",
                    "file_size": fs, "operation_time": _time.perf_counter() - t0}

        tasks = [
            {"index": i, "total": len(pdf_files),
             "input_path": fp,
             "output_path": os.path.join(_outdir, f"{Path(fp).stem}.docx")}
            for i, fp in enumerate(pdf_files)
        ]
        from conversion_worker import ConversionWorker
        self._pdf_to_word_worker = ConversionWorker(tasks, _run_pdf_to_word)
        self._pdf_to_word_worker.file_done.connect(self._on_pdf_to_word_file_done)

        def _on_pdf_to_word_finished(s):
            self.show_progress(False)
            self._set_ui_enabled(True)
            if s['success_count'] > 0 and self.config.get("enable_system_notifications", True):
                self.system_notifier.send("pdf_to_word")
            QMessageBox.information(self, self.translate_text("Succès"),
                self.translate_text("pdf_to_word_success_sum").format(s['success_count'], s['total'], f"{s['total_time']:.1f}"))

        self._pdf_to_word_worker.finished.connect(_on_pdf_to_word_finished)
        self._pdf_to_word_worker.error.connect(lambda e: (
            self.show_progress(False),
            self._set_ui_enabled(True),
            QMessageBox.critical(self, self.translate_text("Erreur"), str(e))
        ))
        self._pdf_to_word_worker.start()

    def _on_pdf_to_word_file_done(self, result):
        if result.get("success"):
            fs = result.get("file_size", 0)
            self.achievement_system.record_conversion("pdf_to_word", fs, True)
            self.achievement_system.mark_format_as_used("pdf")
            self.achievement_system.mark_format_as_used("docx")