"""WordToPdfMixin — Word/DOCX → PDF conversion methods."""

import os
import io
import re as _re
from pathlib import Path
from datetime import datetime

from PySide6.QtWidgets import QMessageBox, QDialog
from PySide6.QtCore import Qt

try:
    from docx import Document
except ImportError as _e:
    Document = None
    print(f"[IMPORT] python-docx not available: {_e}")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import simpleSplit
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
except ImportError as _e:
    letter = A4 = canvas = simpleSplit = None
    getSampleStyleSheet = ParagraphStyle = None
    SimpleDocTemplate = Paragraph = Spacer = None
    print(f"[IMPORT] reportlab not available: {_e}")

from dialogs import WordToPdfOptionsDialog
from conversion_worker import ConversionWorker


def _sanitize_xml(t: str) -> str:
    """Remove control characters that break ReportLab's XML parser."""
    return _re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', t or '')


class WordToPdfMixin:
    """Mixin: Word/DOCX → PDF conversion for FileConverterApp."""

    def convert_word_to_pdf(self):
        if not (hasattr(self, 'active_templates') and 'word_to_pdf' in self.active_templates):
            _def_id, _ = (self._ensure_template_manager() or object()).get_default_template('Conversion Word→PDF')
            if _def_id:
                (self._ensure_template_manager() or object()).apply_template(_def_id, self)

        selected_items = self.files_list_widget.selectedItems()
        files_to_process = []

        if selected_items:
            for i in range(self.files_list_widget.count()):
                item = self.files_list_widget.item(i)
                if item.isSelected():
                    files_to_process.append(item.data(Qt.UserRole))
        else:
            files_to_process = [f for f in self.files_list if f.lower().endswith(('.docx', '.doc'))]

        word_files = [f for f in files_to_process if f.lower().endswith(('.docx', '.doc'))]

        if not word_files:
            msg = self.translate_text("Veuillez sélectionner au moins un fichier Word") if selected_items else self.translate_text("La liste doit contenir au moins un fichier Word")
            QMessageBox.warning(self, self.translate_text("Avertissement"), msg)
            return

        has_formatted_content = False
        for word_file in word_files[:1]:
            try:
                doc = Document(word_file)
                if len(doc.inline_shapes) > 0:
                    has_formatted_content = True
                    break
                if len(doc.tables) > 0:
                    has_formatted_content = True
                    break
                for paragraph in doc.paragraphs[:10]:
                    if paragraph.style.name not in ['Normal', 'Default Paragraph Font']:
                        has_formatted_content = True
                        break
            except Exception:
                pass

        if hasattr(self, 'active_templates') and 'word_to_pdf' in self.active_templates:
            _tpl = self.active_templates['word_to_pdf']
            _quality_map = {
                "Haute (300 DPI)":    "Haute qualité (300 DPI)",
                "Standard (150 DPI)": "Qualité standard (150 DPI)",
                "Basse (72 DPI)":     "Optimisé (96 DPI)",
            }
            options = {
                'mode':             _tpl.get('mode', 'preserve_all'),
                'quality':          _quality_map.get(_tpl.get('quality', ''), 'Qualité standard (150 DPI)'),
                'compress_images':  _tpl.get('compress_images', True),
                'include_metadata': _tpl.get('include_metadata', True),
            }
        else:
            dialog = WordToPdfOptionsDialog(self, self.current_language, has_formatted_content)
            if self.config.get('word_to_pdf_mode') == 'text_only':
                dialog.text_only_radio.setChecked(True)
            if dialog.exec() != QDialog.Accepted:
                return
            options = dialog.get_conversion_mode()

        output_dir = self.get_output_directory()
        if not output_dir:
            return

        message = self.translate_text("conversion_word_to_pdf").format(len(word_files))
        self.show_progress(True, message)
        self._set_ui_enabled(False)

        _options = options
        _outdir  = output_dir

        def _run_word_to_pdf(task):
            import os, time as _time
            t0  = _time.perf_counter()
            fp  = task["input_path"]
            out = task["output_path"]
            fs  = os.path.getsize(fp) if os.path.exists(fp) else 0
            if _options['mode'] == 'preserve_all':
                self.convert_docx_to_pdf_preserve_all(fp, out, _options)
            else:
                self.convert_docx_to_pdf_text_only(fp, out)
            return {"success": True, "error": "",
                    "file_size": fs, "operation_time": _time.perf_counter() - t0}

        tasks = [
            {"index": i, "total": len(word_files),
             "input_path": fp,
             "output_path": os.path.join(_outdir, f"{Path(fp).stem}.pdf")}
            for i, fp in enumerate(word_files)
        ]

        def _on_file_done(result):
            if result.get("success"):
                fp  = result.get("input_path", "?")
                out = result.get("output_path", "?")
                fs  = result.get("file_size", 0)
                ot  = result.get("operation_time", 0)
                print(f"[ACHIEVEMENTS] word_to_pdf file_done: {fp} -> {out}, size={fs}")
                self.achievement_system.record_conversion("word_to_pdf", fs, True)
                self.achievement_system.mark_format_as_used("docx")
                self.achievement_system.mark_format_as_used("pdf")
                if 0 <= datetime.now().hour < 6:
                    self.achievement_system.increment_stat("night_conversions", 1)
                self.db_manager.add_conversion_record(
                    source_file=fp, source_format="DOCX",
                    target_file=out, target_format="PDF",
                    operation_type="word_to_pdf", file_size=fs,
                    conversion_time=ot, success=True,
                    notes=f"Mode: {_options['mode']}"
                )
            else:
                fp = result["input_path"]
                self.db_manager.add_conversion_record(
                    source_file=fp, source_format="DOCX",
                    target_file="", target_format="PDF",
                    operation_type="word_to_pdf", file_size=0,
                    conversion_time=0, success=False,
                    notes=f"Error: {result.get('error','')}"
                )

        def _on_finished(summary):
            self.show_progress(False)
            self._set_ui_enabled(True)
            sc         = summary["success_count"]
            total      = summary["total"]
            total_time = summary["total_time"]
            self.achievement_system.check_achievement("night_owl")
            mode_name = (self.translate_text("Conserver toute la mise en forme")
                         if _options['mode'] == 'preserve_all'
                         else self.translate_text("Texte seulement"))
            formatted_time = f"{total_time:.1f}"
            msg = self.translate_text("word_to_pdf_success_sum").format(
                sc, total, formatted_time, mode_name)
            if self.config.get("enable_system_notifications", True):
                self.system_notifier.send("word_to_pdf")
            if sc >= 50 and total_time <= 300:
                self.achievement_system.update_stat("recent_batch_files", sc)
                self.achievement_system.update_stat("recent_batch_time", total_time)
                self.achievement_system.check_speed_conversion(sc, total_time)
            QMessageBox.information(self, self.translate_text("Succès"),
                                    self.translate_text(msg))

        self._worker = ConversionWorker(tasks, _run_word_to_pdf)
        self._worker.progress.connect(self.progress_bar.setValue)
        self._worker.file_done.connect(_on_file_done)
        self._worker.finished.connect(_on_finished)
        self._worker.start()

    def convert_word_to_pdf_com(self, input_path, output_path, progress_callback=None):
        import os
        import pythoncom

        input_path = os.path.abspath(input_path)
        output_path = os.path.abspath(output_path)

        word = None
        doc = None

        if progress_callback:
            progress_callback(40)

        try:
            pythoncom.CoInitialize()
            import comtypes.client

            if progress_callback:
                progress_callback(50)

            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0

            if progress_callback:
                progress_callback(60)

            doc = word.Documents.Open(input_path)

            if progress_callback:
                progress_callback(80)

            doc.SaveAs2(output_path, FileFormat=17)

            if progress_callback:
                progress_callback(95)

            print(f"[SUCCESS] COM conversion successful: {output_path}")
            return True

        except Exception as e:
            print(f"[ERROR] COM conversion failed: {e}")
            import traceback
            traceback.print_exc()
            return False

        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def convert_docx_to_pdf_preserve_all(self, docx_path, pdf_path, options, progress_callback=None):
        if progress_callback:
            progress_callback(10)

        try:
            if self.convert_word_to_pdf_com(docx_path, pdf_path, progress_callback):
                return True
        except Exception as e:
            print(f"[Word→PDF] COM failed: {e}")

        if progress_callback:
            progress_callback(30)

        try:
            if self._convert_docx_to_pdf_libreoffice(docx_path, pdf_path, progress_callback):
                return True
        except Exception as e:
            print(f"[Word→PDF] LibreOffice failed: {e}")

        if progress_callback:
            progress_callback(50)

        try:
            return self._convert_docx_to_pdf_fallback_reportlab(docx_path, pdf_path, options, progress_callback)
        except Exception as e:
            print(f"[Word→PDF] ReportLab fallback failed: {e}")
            return False

    def _convert_docx_to_pdf_libreoffice(self, docx_path, pdf_path, progress_callback=None):
        import subprocess
        import shutil

        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files\The Document Foundation\LibreOffice\program\soffice.exe",
        ]

        soffice = None
        for path in libreoffice_paths:
            if os.path.exists(path):
                soffice = path
                break

        if not soffice:
            soffice = shutil.which("soffice")

        if not soffice:
            print("[Word→PDF] LibreOffice not found")
            return False

        if progress_callback:
            progress_callback(40)

        try:
            output_dir = os.path.dirname(pdf_path)
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", output_dir, docx_path],
                capture_output=True, text=True, timeout=120,
                creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0)
            )

            if progress_callback:
                progress_callback(90)

            expected_pdf = os.path.join(output_dir,
                os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")

            if os.path.exists(expected_pdf):
                if expected_pdf != pdf_path:
                    os.replace(expected_pdf, pdf_path)
                print(f"[Word→PDF] LibreOffice success: {pdf_path}")
                return True

            print(f"[Word→PDF] LibreOffice failed: {result.stderr}")
            return False

        except subprocess.TimeoutExpired:
            print("[Word→PDF] LibreOffice timed out")
            return False
        except Exception as e:
            print(f"[Word→PDF] LibreOffice error: {e}")
            return False

    def _convert_docx_to_pdf_fallback_reportlab(self, docx_path, pdf_path, options, progress_callback=None):
        if canvas is None:
            print("[Word→PDF] ReportLab not available")
            self.create_empty_pdf_with_message(pdf_path, "ReportLab not available")
            return False

        try:
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 72

            if progress_callback:
                progress_callback(60)

            def ensure_space(needed):
                nonlocal y
                if y - needed < 72:
                    c.showPage()
                    y = height - 72

            def wrap_text(text, font_name, font_size, max_w):
                c.setFont(font_name, font_size)
                return simpleSplit(text, font_name, font_size, max_w)

            def draw_paragraph(para):
                nonlocal y
                text = para.text.strip()
                if not text:
                    y -= 12
                    return

                style_name = para.style.name if para.style else "Normal"

                if 'Heading 1' in style_name or 'Titre 1' in style_name:
                    font_name = 'Helvetica-Bold'
                    font_size = 16
                    spacing = 20
                elif 'Heading 2' in style_name or 'Titre 2' in style_name:
                    font_name = 'Helvetica-Bold'
                    font_size = 14
                    spacing = 16
                elif 'Heading 3' in style_name or 'Titre 3' in style_name:
                    font_name = 'Helvetica-Bold'
                    font_size = 12
                    spacing = 14
                elif 'List' in style_name:
                    font_name = 'Helvetica'
                    font_size = 11
                    spacing = 12
                else:
                    font_name = 'Helvetica'
                    font_size = 11
                    spacing = 14

                lines = wrap_text(text, font_name, font_size, width - 144)
                total_height = len(lines) * spacing

                ensure_space(total_height)

                for line in lines:
                    c.setFont(font_name, font_size)
                    c.drawString(72, y, line)
                    y -= spacing

                y -= 4

            def draw_image_from_blob(blob):
                nonlocal y
                try:
                    from PIL import Image as PILImage
                    import tempfile

                    img = PILImage.open(io.BytesIO(blob))
                    img_w, img_h = img.size

                    max_w = width - 144
                    max_h = 200

                    scale = min(max_w / img_w, max_h / img_h, 1.0)
                    new_w = img_w * scale
                    new_h = img_h * scale

                    ensure_space(new_h + 20)

                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        img.save(tmp.name, 'PNG')
                        c.drawImage(tmp.name, 72, y - new_h, new_w, new_h)
                        os.unlink(tmp.name)

                    y -= new_h + 20

                except Exception as e:
                    print(f"[Image render error] {e}")

            def draw_table(tbl):
                nonlocal y
                try:
                    rows = tbl.rows
                    n_rows = len(rows)
                    n_cols = max(len(row.cells) for row in rows)

                    if n_cols == 0:
                        return

                    col_width = (width - 144) / n_cols
                    row_height = 20

                    total_height = n_rows * row_height
                    ensure_space(total_height + 10)

                    for row in rows:
                        for i, cell in enumerate(row.cells):
                            x = 72 + i * col_width
                            c.rect(x, y - row_height, col_width, row_height)
                            text = cell.text.strip()[:50]
                            c.setFont('Helvetica', 8)
                            c.drawString(x + 2, y - 14, text)
                        y -= row_height

                    y -= 10

                except Exception as e:
                    print(f"[Table render error] {e}")

            def _get_inline_image_blob(inline_elem):
                try:
                    blip = inline_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}blip')
                    if blip is not None:
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId:
                            rel = doc.part.rels.get(rId)
                            if rel and hasattr(rel, 'target_part'):
                                return rel.target_part.blob
                except Exception:
                    pass
                return None

            for para in doc.paragraphs:
                draw_paragraph(para)

                for run in para.runs:
                    if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        try:
                            drawing = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                            if drawing is not None:
                                blob = _get_inline_image_blob(drawing)
                                if blob:
                                    draw_image_from_blob(blob)
                        except Exception:
                            pass

            for table in doc.tables:
                draw_table(table)

            if progress_callback:
                progress_callback(90)

            c.save()
            print(f"[Word→PDF] ReportLab fallback success: {pdf_path}")
            return True

        except Exception as e:
            print(f"[Word→PDF] ReportLab error: {e}")
            return False

    def convert_docx_to_pdf_text_only(self, docx_path, pdf_path):
        try:
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT

            print(f"[DEBUG] Starting text-only conversion: {docx_path}")

            doc = Document(docx_path)

            buffer = io.BytesIO()

            doc_template = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            story = []
            styles = getSampleStyleSheet()

            style_title = ParagraphStyle(
                'TitleStyle',
                parent=styles['Heading1'],
                fontSize=18,
                leading=22,
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                textColor='#2c3e50'
            )

            style_heading = ParagraphStyle(
                'HeadingStyle',
                parent=styles['Heading2'],
                fontSize=14,
                leading=18,
                spaceAfter=10,
                spaceBefore=15,
                fontName='Helvetica-Bold',
                textColor='#34495e'
            )

            style_normal = ParagraphStyle(
                'CleanNormal',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                spaceAfter=8,
                alignment=TA_JUSTIFY,
                wordWrap='LTR'
            )

            style_bullet = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                spaceAfter=6,
                leftIndent=20,
                firstLineIndent=-15
            )

            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    style_name = str(paragraph.style.name) if paragraph.style else "Normal"

                    if i == 0 and len(text) < 100:
                        style = style_title
                    elif 'Heading' in style_name or 'Titre' in style_name:
                        style = style_heading
                    elif text.startswith('•') or text.startswith('-') or text.startswith('*'):
                        style = style_bullet
                        text = "• " + text.lstrip('•-* ')
                    else:
                        style = style_normal

                    if hasattr(paragraph, 'alignment'):
                        if paragraph.alignment == 1:
                            style = ParagraphStyle(
                                'CenterText',
                                parent=style,
                                alignment=TA_CENTER
                            )
                        elif paragraph.alignment == 2:
                            style = ParagraphStyle(
                                'RightText',
                                parent=style,
                                alignment=TA_RIGHT
                            )

                    p = Paragraph(_sanitize_xml(text), style)
                    story.append(p)
                    story.append(Spacer(1, 6))

            if story:
                story.append(Spacer(1, 20))
                note_style = ParagraphStyle(
                    'NoteStyle',
                    parent=styles['Normal'],
                    fontSize=9,
                    leading=12,
                    alignment=TA_CENTER,
                    textColor='#7f8c8d',
                    fontName='Helvetica-Oblique'
                )
                note_text = "Note : Version texte seulement - Les images et tableaux ont été omis pour une meilleure lisibilité"
                story.append(Paragraph(note_text, note_style))

            if story:
                doc_template.build(story)
            else:
                empty_style = ParagraphStyle(
                    'EmptyStyle',
                    parent=styles['Normal'],
                    fontSize=12,
                    leading=16,
                    alignment=TA_CENTER,
                    textColor='#95a5a6'
                )
                story.append(Paragraph("Document vide", empty_style))
                doc_template.build(story)

            with open(pdf_path, 'wb') as f:
                f.write(buffer.getvalue())

            buffer.close()

            print(f"[DEBUG] Text-only conversion complete: {pdf_path}")

        except Exception as e:
            print(f"[ERROR] Text-only conversion error: {e}")
            self.convert_docx_to_pdf_simple(docx_path, pdf_path)

    def analyze_word_content(self, docx_path):
        try:
            doc = Document(docx_path)

            content_info = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'images': 0,
                'has_complex_layout': False,
                'total_text_length': 0,
            }

            for para in doc.paragraphs:
                content_info['total_text_length'] += len(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content_info['total_text_length'] += len(cell.text)

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    content_info['images'] += 1

            if content_info['tables'] > 3:
                content_info['has_complex_layout'] = True
            if content_info['images'] > 5:
                content_info['has_complex_layout'] = True

            return content_info

        except Exception as e:
            print(f"[Analyze] Error: {e}")
            return {
                'paragraphs': 0, 'tables': 0, 'images': 0,
                'has_complex_layout': False, 'total_text_length': 0
            }

    def convert_docx_to_pdf_advanced(self, docx_path, pdf_path):
        try:
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 72

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    y -= 12
                    continue

                style_name = para.style.name if para.style else "Normal"

                if 'Heading 1' in style_name:
                    font_name = 'Helvetica-Bold'
                    font_size = 16
                    spacing = 20
                elif 'Heading 2' in style_name:
                    font_name = 'Helvetica-Bold'
                    font_size = 14
                    spacing = 16
                elif 'Heading 3' in style_name:
                    font_name = 'Helvetica-Bold'
                    font_size = 12
                    spacing = 14
                else:
                    font_name = 'Helvetica'
                    font_size = 11
                    spacing = 14

                lines = simpleSplit(text, font_name, font_size, width - 144)
                total_height = len(lines) * spacing

                if y - total_height < 72:
                    c.showPage()
                    y = height - 72

                for line in lines:
                    c.setFont(font_name, font_size)
                    c.drawString(72, y, line)
                    y -= spacing

                y -= 4

            c.save()
            return True

        except Exception as e:
            print(f"[Word→PDF] Advanced error: {e}")
            return False

    def convert_docx_to_pdf_simple(self, docx_path, pdf_path):
        try:
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 72

            c.setFont('Helvetica-Bold', 16)
            title = Path(docx_path).stem
            c.drawString(72, y, title[:60])
            y -= 30

            c.setFont('Helvetica', 11)
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    y -= 12
                    continue

                lines = simpleSplit(text, 'Helvetica', 11, width - 144)

                if y - len(lines) * 14 < 72:
                    c.showPage()
                    y = height - 72

                for line in lines:
                    c.drawString(72, y, line)
                    y -= 14

                y -= 4

            c.save()
            return True

        except Exception as e:
            print(f"[Word→PDF] Simple error: {e}")
            return False

    def extract_table_from_element(self, element, doc):
        try:
            from docx.oxml.ns import qn

            tbl = element.find(qn('w:tbl'))
            if tbl is None:
                return None

            rows = []
            for tr in tbl.findall(qn('w:tr')):
                cells = []
                for tc in tr.findall(qn('w:tc')):
                    text = ''
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            for t in r.findall(qn('w:t')):
                                if t.text:
                                    text += t.text
                    cells.append(text)
                rows.append(cells)

            return rows

        except Exception as e:
            print(f"[Extract table] Error: {e}")
            return None

    def convert_docx_to_pdf_with_images(self, docx_path, pdf_path):
        try:
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 72

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    y -= 12
                    continue

                lines = simpleSplit(text, 'Helvetica', 11, width - 144)

                if y - len(lines) * 14 < 72:
                    c.showPage()
                    y = height - 72

                for line in lines:
                    c.setFont('Helvetica', 11)
                    c.drawString(72, y, line)
                    y -= 14

                y -= 4

            for table in doc.tables:
                try:
                    for row in table.rows:
                        if y < 100:
                            c.showPage()
                            y = height - 72

                        x = 72
                        for cell in row.cells:
                            c.rect(x, y - 20, 80, 20)
                            c.setFont('Helvetica', 8)
                            c.drawString(x + 2, y - 14, cell.text.strip()[:30])
                            x += 82
                        y -= 24
                except Exception:
                    pass

            c.save()
            return True

        except Exception as e:
            print(f"[Word→PDF] With images error: {e}")
            return False

    def convert_docx_to_pdf_fallback(self, docx_path, pdf_path):
        try:
            return self.convert_docx_to_pdf_simple(docx_path, pdf_path)
        except Exception as e:
            print(f"[Word→PDF] Fallback error: {e}")
            self.create_empty_pdf_with_message(pdf_path, f"Erreur: {str(e)[:100]}")
            return False

    def create_minimal_pdf_from_docx(self, docx_path, pdf_path):
        try:
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4

            c.setFont('Helvetica-Bold', 14)
            c.drawString(72, height - 72, Path(docx_path).stem[:50])

            c.setFont('Helvetica', 11)
            y = height - 100
            for para in doc.paragraphs[:50]:
                text = para.text.strip()
                if text:
                    lines = simpleSplit(text, 'Helvetica', 11, width - 144)
                    for line in lines[:3]:
                        if y < 72:
                            break
                        c.drawString(72, y, line)
                        y -= 14
                    y -= 6

            c.save()
            return True

        except Exception as e:
            print(f"[Word→PDF] Minimal error: {e}")
            return False

    def create_empty_pdf_with_message(self, pdf_path, message):
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica", 12)
        c.drawString(100, height - 150, message)
        c.drawString(100, height - 170, "The original document could not be converted correctly.")

        c.save()

    def _open_image_for_pdf(self, file_path):
        from PIL import Image
        return Image.open(file_path)