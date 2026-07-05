"""PreviewDialog — Quick file preview (PDF, Images, Word, Text, CSV, JSON, XLSX, HTML, EPUB, Audio, Video)."""

import os
import tempfile
from pathlib import Path

import fitz
from PySide6.QtWidgets import (QDialog, QVBoxLayout, QLabel, QScrollArea,
                               QWidget, QPushButton, QTableWidget,
                               QTableWidgetItem, QHeaderView, QTabWidget)
from PySide6.QtCore import Qt
from PySide6.QtGui import QPixmap

from utils import make_tm
from utils.translation_mixin import TranslationMixin


class PreviewDialog(TranslationMixin, QDialog):
    def __init__(self, file_path, parent=None, language="fr"):
        super().__init__(parent)
        self.file_path = file_path
        self.language = language
        self._tm = make_tm(language)
        self._media_player = None  # QMediaPlayer ref kept alive
        self.setWindowTitle(self.translate_text("Aperçu -") + f" {Path(file_path).name}")
        self.setModal(False)
        self.setMinimumSize(700, 600)
        self.resize(800, 680)
        self.setup_ui()
        self.load_preview()

    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        self.header_label = QLabel(f"<b>{Path(self.file_path).name}</b>")
        self.header_label.setStyleSheet("font-size: 14px; padding: 4px 8px;")
        layout.addWidget(self.header_label)

        self.content_area = QWidget()
        self.content_layout = QVBoxLayout(self.content_area)
        self.content_layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.content_area, 1)

        self.preview_label = QLabel()
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setWordWrap(True)
        self.content_layout.addWidget(self.preview_label)

        self.close_btn = QPushButton(self.translate_text("Fermer l'aperçu"))
        self.close_btn.setFixedHeight(36)
        self.close_btn.clicked.connect(self.close)
        layout.addWidget(self.close_btn)

    def _clear_content(self):
        """Remove all widgets from content_layout (except preview_label)."""
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            w = item.widget()
            if w and w is not self.preview_label:
                w.deleteLater()
        self.preview_label.clear()
        self.preview_label.setParent(None)

    def _show_error(self, msg):
        self.preview_label.setParent(self.content_area)
        self.content_layout.addWidget(self.preview_label)
        self.preview_label.setText(f"⚠️ {msg}")

    def load_preview(self):
        ext = Path(self.file_path).suffix.lower()
        try:
            if ext == '.pdf':
                self.preview_pdf()
            elif ext in (
                '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif',
                '.svg', '.ico', '.avif', '.heic', '.heif', '.webp', '.jfif',
                '.psd', '.j2k', '.jp2', '.jpx', '.dng', '.cr2',
                '.cr3', '.nef', '.arw', '.orf', '.rw2', '.raf',
            ):
                self.preview_image()
            elif ext in ('.docx', '.doc'):
                self.preview_word()
            elif ext in ('.txt', '.log', '.md', '.py', '.js', '.ts', '.css', '.xml', '.yaml', '.yml', '.ini', '.cfg', '.bat', '.sh'):
                self.preview_text(syntax=ext)
            elif ext == '.rtf':
                self.preview_rtf()
            elif ext == '.csv':
                self.preview_csv()
            elif ext == '.json':
                self.preview_json()
            elif ext in ('.xlsx', '.xls'):
                self.preview_xlsx()
            elif ext in ('.html', '.htm'):
                self.preview_html()
            elif ext == '.epub':
                self.preview_epub()
            elif ext in ('.mp3', '.wav', '.flac', '.ogg', '.m4a', '.aac', '.wma'):
                self.preview_audio()
            elif ext in ('.mp4', '.avi', '.mkv', '.mov', '.webm', '.wmv', '.flv'):
                self.preview_video()
            else:
                self.preview_unsupported()
        except Exception as e:
            translated_msg = self.translate_text('Erreur lors du chargement de l\'aperçu:')
            self._show_error(f"{translated_msg}\n{e}")

    def preview_pdf(self):
        try:
            pdf_document = fitz.open(self.file_path)
            if len(pdf_document) == 0:
                self._show_error(self.translate_text("PDF vide"))
                return

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            container = QWidget()
            vbox = QVBoxLayout(container)
            vbox.setSpacing(8)

            max_pages = min(len(pdf_document), 5)
            for i in range(max_pages):
                page = pdf_document.load_page(i)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                tmp.close()
                pix.save(tmp.name)
                pixmap = QPixmap(tmp.name)
                os.unlink(tmp.name)
                if not pixmap.isNull():
                    if pixmap.width() > 660:
                        pixmap = pixmap.scaledToWidth(660, Qt.SmoothTransformation)
                    lbl = QLabel()
                    lbl.setPixmap(pixmap)
                    lbl.setAlignment(Qt.AlignHCenter)
                    vbox.addWidget(lbl)

            if len(pdf_document) > max_pages:
                more = QLabel(f"… {len(pdf_document) - max_pages} {self.translate_text('pages supplémentaires')}")
                more.setAlignment(Qt.AlignCenter)
                more.setStyleSheet("color: gray; font-style: italic; padding: 6px;")
                vbox.addWidget(more)

            pdf_document.close()
            vbox.addStretch()
            scroll.setWidget(container)
            self._clear_content()
            self.content_layout.addWidget(scroll)

        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur PDF:')} {e}")

    _QPIXMAP_NATIVE = frozenset({
        ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".ico",
    })

    def _load_pixmap(self, path: str) -> "QPixmap":
        """
        Load any image as a QPixmap.
        Priority:
          1. QPixmap directly          — fast, for Qt-native formats
          2. SVG via QSvgRenderer      — vector, lossless
          3. pillow-heif               — HEIC/HEIF
          4. rawpy                     — DNG/CR2/CR3/NEF/ARW/ORF/RW2/RAF
          5. psd-tools                 — PSD (flatten layers)
          6. Pillow generic            — AVIF, TIFF, J2K, JP2, WEBP anim, etc.
          7. QPixmap fallback          — last resort (returns null if unsupported)
        """
        from PySide6.QtGui import QImage
        ext = Path(path).suffix.lower()

        if ext in self._QPIXMAP_NATIVE:
            px = QPixmap(path)
            if not px.isNull():
                return px

        if ext == ".svg":
            try:
                from PySide6.QtSvg import QSvgRenderer
                from PySide6.QtGui import QPainter
                renderer = QSvgRenderer(path)
                sz = renderer.defaultSize()
                if not sz.isValid():
                    sz = renderer.viewBox().size()
                w = max(sz.width(),  1)
                h = max(sz.height(), 1)
                img = QImage(w, h, QImage.Format_ARGB32)
                img.fill(0)
                painter = QPainter(img)
                renderer.render(painter)
                painter.end()
                return QPixmap.fromImage(img)
            except Exception:
                pass

        if ext in (".heic", ".heif"):
            try:
                from pillow_heif import register_heif_opener
                register_heif_opener()
                from PIL import Image
                img = Image.open(path).convert("RGBA")
                data = img.tobytes("raw", "RGBA")
                qi = QImage(data, img.width, img.height, QImage.Format_RGBA8888)
                return QPixmap.fromImage(qi)
            except Exception:
                pass

        if ext in (".dng", ".cr2", ".cr3", ".nef", ".arw", ".orf", ".rw2", ".raf"):
            try:
                import rawpy
                with rawpy.imread(path) as raw:
                    rgb = raw.postprocess(use_camera_wb=True, output_bps=8)
                h, w, _ = rgb.shape
                qi = QImage(rgb.tobytes(), w, h, w * 3, QImage.Format_RGB888)
                return QPixmap.fromImage(qi)
            except Exception:
                pass

        if ext == ".psd":
            try:
                from psd_tools import PSDImage
                psd = PSDImage.open(path)
                img = psd.composite().convert("RGBA")
                data = img.tobytes("raw", "RGBA")
                qi = QImage(data, img.width, img.height, QImage.Format_RGBA8888)
                return QPixmap.fromImage(qi)
            except Exception:
                pass

        try:
            from PIL import Image
            img = Image.open(path)
            try:
                img.seek(0)
            except (AttributeError, EOFError):
                pass
            img = img.convert("RGBA")
            data = img.tobytes("raw", "RGBA")
            qi = QImage(data, img.width, img.height, QImage.Format_RGBA8888)
            px = QPixmap.fromImage(qi)
            if not px.isNull():
                return px
        except Exception:
            pass

        return QPixmap(path)

    def preview_image(self):
        try:
            pixmap = self._load_pixmap(self.file_path)
            if pixmap.isNull():
                self._show_error(self.translate_text("Format d'image non supporté"))
                return
            pixmap = pixmap.scaled(660, 580, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            lbl = QLabel()
            lbl.setPixmap(pixmap)
            lbl.setAlignment(Qt.AlignCenter)
            scroll.setWidget(lbl)
            self._clear_content()
            self.content_layout.addWidget(scroll)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur image:')} {e}")

    def preview_word(self):
        try:
            from docx import Document as _DocxDocument
            doc = _DocxDocument(self.file_path)
            from PySide6.QtWidgets import QTextEdit
            te = QTextEdit()
            te.setReadOnly(True)
            html_parts = []
            for para in doc.paragraphs[:80]:
                text = para.text
                if not text.strip():
                    html_parts.append("<br>")
                    continue
                style = para.style.name if para.style else ""
                if "Heading 1" in style:
                    html_parts.append(f"<h2>{text}</h2>")
                elif "Heading 2" in style:
                    html_parts.append(f"<h3>{text}</h3>")
                else:
                    html_parts.append(f"<p style='margin:2px 0'>{text}</p>")
            if len(doc.paragraphs) > 80:
                html_parts.append("<p style='color:gray;font-style:italic'>[...]</p>")
            te.setHtml("".join(html_parts) if html_parts else self.translate_text("Document vide"))
            self._clear_content()
            self.content_layout.addWidget(te)
        except Exception as e:
            self._show_error(f"{self.translate_text('Impossible de prévisualiser le document Word')}\n{e}")

    def preview_text(self, syntax=""):
        try:
            from PySide6.QtWidgets import QTextEdit
            from PySide6.QtGui import QFont
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read(200_000)
            te = QTextEdit()
            te.setReadOnly(True)
            te.setFont(QFont("Consolas, Courier New, monospace", 10))
            te.setPlainText(content)
            self._clear_content()
            self.content_layout.addWidget(te)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur texte:')} {e}")

    def preview_rtf(self):
        name = Path(self.file_path).name
        self.preview_label.setParent(self.content_area)
        self.content_layout.addWidget(self.preview_label)
        self.preview_label.setText(
            f"🚫  {self.translate_text('Aperçu non disponible pour ce type de fichier')}\n({name})"
        )

    def preview_csv(self):
        try:
            import csv
            with open(self.file_path, 'r', encoding='utf-8', errors='replace', newline='') as f:
                reader = csv.reader(f)
                rows = [r for _, r in zip(range(200), reader)]

            if not rows:
                self._show_error(self.translate_text("Fichier CSV vide"))
                return

            ncols = max(len(r) for r in rows)
            table = QTableWidget(len(rows), ncols)
            table.setHorizontalHeader
            header_row = rows[0]
            table.setHorizontalHeaderLabels(header_row + [''] * (ncols - len(header_row)))
            for ri, row in enumerate(rows[1:], 0):
                for ci, val in enumerate(row):
                    table.setItem(ri, ci, QTableWidgetItem(val))
            table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
            table.horizontalHeader().setStretchLastSection(True)
            table.setEditTriggers(QTableWidget.NoEditTriggers)
            self._clear_content()
            self.content_layout.addWidget(table)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur CSV:')} {e}")

    def preview_json(self):
        try:
            import json
            from PySide6.QtWidgets import QTextEdit
            from PySide6.QtGui import QFont, QColor, QTextCharFormat, QSyntaxHighlighter
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                raw = f.read(500_000)
            try:
                parsed = json.loads(raw)
                pretty = json.dumps(parsed, indent=2, ensure_ascii=False)
            except Exception:
                pretty = raw

            te = QTextEdit()
            te.setReadOnly(True)
            te.setFont(QFont("Consolas", 10))
            te.setPlainText(pretty)

            class _JsonHL(QSyntaxHighlighter):
                def highlightBlock(self, text):
                    import re
                    fmt_key   = QTextCharFormat(); fmt_key.setForeground(QColor("#569cd6"))
                    fmt_str   = QTextCharFormat(); fmt_str.setForeground(QColor("#ce9178"))
                    fmt_num   = QTextCharFormat(); fmt_num.setForeground(QColor("#b5cea8"))
                    fmt_kw    = QTextCharFormat(); fmt_kw.setForeground(QColor("#569cd6"))
                    for m in re.finditer(r'"[^"\\]*(?:\\.[^"\\]*)*"\s*:', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_key)
                    for m in re.finditer(r'(?<!\w)"[^"\\]*(?:\\.[^"\\]*)*"', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_str)
                    for m in re.finditer(r'\b-?\d+\.?\d*([eE][+-]?\d+)?\b', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_num)
                    for m in re.finditer(r'\b(true|false|null)\b', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_kw)

            _JsonHL(te.document())
            self._clear_content()
            self.content_layout.addWidget(te)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur JSON:')} {e}")

    def preview_xlsx(self):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(self.file_path, read_only=True, data_only=True)
            tabs = QTabWidget()
            for sheet_name in wb.sheetnames[:6]:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(max_row=200, values_only=True))
                if not rows:
                    continue
                ncols = max((len(r) for r in rows), default=0)
                table = QTableWidget(len(rows), ncols)
                table.setHorizontalHeaderLabels([str(i+1) for i in range(ncols)])
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        table.setItem(ri, ci, QTableWidgetItem(str(val) if val is not None else ""))
                table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
                table.horizontalHeader().setStretchLastSection(True)
                table.setEditTriggers(QTableWidget.NoEditTriggers)
                tabs.addTab(table, sheet_name)
            wb.close()
            self._clear_content()
            self.content_layout.addWidget(tabs)
        except ImportError:
            self._show_error("openpyxl non installé — pip install openpyxl")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur XLSX:')} {e}")

    def preview_html(self):
        try:
            from PySide6.QtWebEngineWidgets import QWebEngineView
            from PySide6.QtCore import QUrl

            view = QWebEngineView()
            view.setUrl(QUrl.fromLocalFile(str(Path(self.file_path).resolve())))
            self._clear_content()
            self.content_layout.addWidget(view)
            self.resize(900, 700)
            return
        except ImportError:
            pass
        except Exception:
            pass

        self._html_fallback_view()

    def _html_fallback_view(self):
        """
        Fallback when QtWebEngineWidgets is missing from the build:
        displays the HTML source code with full syntax highlighting
        (HTML + inline CSS in <style> + inline JS in <script>)
        and an explicit warning banner for the user.
        """
        from PySide6.QtWidgets import QTextEdit, QLabel
        from PySide6.QtGui import QFont, QColor, QTextCharFormat, QSyntaxHighlighter
        import re

        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                source = f.read(500_000)
        except Exception as e:
            self._show_error(f"Erreur lecture HTML : {e}")
            return

        te = QTextEdit()
        te.setReadOnly(True)
        te.setFont(QFont("Consolas", 10))

        class _HtmlCssJsHL(QSyntaxHighlighter):
            """
            Three-level syntax highlighting:
              • HTML  — tags, attributes, values, comments, DOCTYPE
              • CSS   — selectors, properties, values, @-rules, comments
                        inside each <style>…</style> block
              • JS    — keywords, strings, numbers, comments, regex
                        inside each <script>…</script> block
            """

            def __init__(self, document, source_text=""):
                super().__init__(document)

                def _fmt(hex_color, bold=False, italic=False):
                    f = QTextCharFormat()
                    f.setForeground(QColor(hex_color))
                    if bold:   f.setFontWeight(700)
                    if italic: f.setFontItalic(True)
                    return f

                # HTML
                self._html_tag     = _fmt("#569cd6", bold=True)
                self._html_attr    = _fmt("#9cdcfe")
                self._html_val     = _fmt("#ce9178")
                self._html_comment = _fmt("#6a9955", italic=True)
                self._html_doctype = _fmt("#c586c0")
                self._html_entity  = _fmt("#f0a070")

                # CSS
                self._css_selector = _fmt("#d7ba7d", bold=True)
                self._css_prop     = _fmt("#9cdcfe")
                self._css_val      = _fmt("#ce9178")
                self._css_atrule   = _fmt("#c586c0", bold=True)
                self._css_comment  = _fmt("#6a9955", italic=True)
                self._css_number   = _fmt("#b5cea8")

                # JS
                self._js_keyword   = _fmt("#569cd6", bold=True)
                self._js_builtin   = _fmt("#4ec9b0")
                self._js_string    = _fmt("#ce9178")
                self._js_template  = _fmt("#ce9178")
                self._js_number    = _fmt("#b5cea8")
                self._js_comment   = _fmt("#6a9955", italic=True)
                self._js_regex     = _fmt("#d16969")
                self._js_operator  = _fmt("#d4d4d4")

                self._JS_KEYWORDS = {
                    "var","let","const","function","class","return","if","else",
                    "for","while","do","switch","case","break","continue","new",
                    "delete","typeof","instanceof","in","of","try","catch",
                    "finally","throw","async","await","import","export","from",
                    "default","extends","super","this","static","yield","void",
                    "null","undefined","true","false","debugger","with",
                }
                self._JS_BUILTINS = {
                    "console","document","window","Array","Object","String",
                    "Number","Boolean","Math","Date","JSON","Promise","fetch",
                    "setTimeout","setInterval","clearTimeout","clearInterval",
                    "parseInt","parseFloat","isNaN","isFinite","encodeURI",
                    "decodeURI","Map","Set","WeakMap","WeakSet","Symbol",
                    "Proxy","Reflect","Error","TypeError","RangeError",
                    "localStorage","sessionStorage","navigator","location",
                    "history","screen","alert","confirm","prompt",
                }

                self._style_ranges  = []
                self._script_ranges = []
                if source_text:
                    self._compute_ranges_from_source(source_text)

            def _compute_ranges_from_source(self, text):
                """
                Computes line ranges for <style> and <script> blocks directly
                from raw text (without touching the QTextDocument).
                Called once in __init__, never from highlightBlock.
                """
                text.splitlines()

                def _line_of_pos(pos):
                    return text[:pos].count("\n")

                for open_pat, close_pat, target in [
                    (r'<style[^>]*>',  r'</style\s*>',  self._style_ranges),
                    (r'<script[^>]*>', r'</script\s*>', self._script_ranges),
                ]:
                    for mo in re.finditer(open_pat, text, re.IGNORECASE):
                        inner_start = mo.end()
                        mc = re.search(close_pat, text[inner_start:], re.IGNORECASE)
                        inner_end = inner_start + mc.start() if mc else len(text)
                        target.append((_line_of_pos(inner_start), _line_of_pos(inner_end)))

            def _in_range(self, block_num, ranges):
                return any(s <= block_num <= e for s, e in ranges)

            def _apply(self, text, pattern, fmt, flags=0):
                for m in re.finditer(pattern, text, flags):
                    self.setFormat(m.start(), len(m.group()), fmt)

            def highlightBlock(self, text):
                bn = self.currentBlock().blockNumber()
                in_style  = self._in_range(bn, self._style_ranges)
                in_script = self._in_range(bn, self._script_ranges)

                if in_style:
                    self._highlight_css(text)
                elif in_script:
                    self._highlight_js(text)
                else:
                    self._highlight_html(text)

            # HTML
            def _highlight_html(self, text):
                self._apply(text, r'<!--.*?-->', self._html_comment)
                self._apply(text, r'<!DOCTYPE[^>]*>', self._html_doctype, re.IGNORECASE)
                self._apply(text, r'</?[\w\-:.]+', self._html_tag)
                self._apply(text, r'\b[\w\-:]+=', self._html_attr)
                for m in re.finditer(r'\b([\w\-:]+)=', text):
                    self.setFormat(m.start(1), len(m.group(1)), self._html_attr)
                self._apply(text, r'(["\'])(?:(?!\1).)*\1', self._html_val)
                self._apply(text, r'&(?:#\d+|#x[\da-fA-F]+|[a-zA-Z]\w*);', self._html_entity)

            # CSS
            def _highlight_css(self, text):
                self._apply(text, r'/\*.*?\*/', self._css_comment)
                self._apply(text, r'@[\w-]+', self._css_atrule)
                self._apply(text, r'[\w-]+(?=\s*:)', self._css_prop)
                self._apply(text, r'-?\d+\.?\d*(?:%|px|em|rem|vh|vw|vmin|vmax|pt|cm|mm|s|ms)?',
                            self._css_number)
                self._apply(text, r'(["\'])(?:(?!\1).)*\1', self._css_val)
                self._apply(text, r'#[0-9a-fA-F]{3,8}\b', self._css_val)
                if re.search(r'\{', text) or (not re.search(r':', text)):
                    self._apply(text, r'[.#]?[\w][\w-]*(?=[^:]*\{)', self._css_selector)

            # JS
            def _highlight_js(self, text):
                self._apply(text, r'//.*', self._js_comment)
                self._apply(text, r'/\*.*?\*/', self._js_comment)
                self._apply(text, r'`(?:[^`\\]|\\.)*`', self._js_template)
                self._apply(text, r'(["\'])(?:(?!\1|\\).|\\.)*\1', self._js_string)
                self._apply(text, r'\b0[xX][\da-fA-F]+\b|\b\d+\.?\d*(?:[eE][+-]?\d+)?\b',
                            self._js_number)
                for m in re.finditer(r'\b([a-zA-Z_$][\w$]*)\b', text):
                    word = m.group(1)
                    if word in self._JS_KEYWORDS:
                        self.setFormat(m.start(), len(word), self._js_keyword)
                    elif word in self._JS_BUILTINS:
                        self.setFormat(m.start(), len(word), self._js_builtin)
                self._apply(text, r'(?<![<\w\d])/(?:[^/\\\n]|\\.)+/[gimsuy]*', self._js_regex)

        _HtmlCssJsHL(te.document(), source)
        te.setPlainText(source)

        warn = QLabel(self.translate_text("preview_unavailable"))
        warn.setStyleSheet(
            "background:#2d2400; color:#ffcc00; "
            "padding:6px 12px; font-size:12px; border-radius:4px;"
        )
        warn.setWordWrap(True)

        self._clear_content()
        self.content_layout.addWidget(warn)
        self.content_layout.addWidget(te)

    def preview_epub(self):
        try:
            import ebooklib
            from ebooklib import epub
            from PySide6.QtWidgets import QTextEdit
            book = epub.read_epub(self.file_path)
            html_parts = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    html_parts.append(item.get_content().decode('utf-8', errors='replace'))
                    if len(html_parts) >= 3:
                        break
            if not html_parts:
                self._show_error(self.translate_text("EPUB vide ou contenu non extractible"))
                return
            combined = "<hr>".join(html_parts[:3])
            te = QTextEdit()
            te.setReadOnly(True)
            te.setHtml(combined)
            self._clear_content()
            self.content_layout.addWidget(te)
        except ImportError:
            self._show_error("ebooklib non installé — pip install ebooklib")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur EPUB:')} {e}")

    def preview_audio(self):
        try:
            from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
            from PySide6.QtCore import QUrl
            from PySide6.QtWidgets import QSlider, QHBoxLayout, QVBoxLayout, QWidget, QLabel, QPushButton

            player = QMediaPlayer()
            audio_out = QAudioOutput()
            player.setAudioOutput(audio_out)
            audio_out.setVolume(0.8)
            player.setSource(QUrl.fromLocalFile(str(Path(self.file_path).resolve())))
            self._media_player = player
            self._audio_output = audio_out

            container = QWidget()
            vbox = QVBoxLayout(container)
            vbox.setSpacing(12)

            info = QLabel(f"🎵  <b>{Path(self.file_path).name}</b>")
            info.setStyleSheet("font-size: 14px; padding: 8px;")
            info.setAlignment(Qt.AlignCenter)
            vbox.addWidget(info)

            time_lbl = QLabel("0:00 / 0:00")
            time_lbl.setAlignment(Qt.AlignCenter)
            time_lbl.setStyleSheet("font-size: 12px; color: gray;")
            vbox.addWidget(time_lbl)

            seek_slider = QSlider(Qt.Horizontal)
            seek_slider.setRange(0, 0)
            vbox.addWidget(seek_slider)

            ctrl = QWidget()
            hbox = QHBoxLayout(ctrl)
            hbox.setSpacing(12)

            play_btn  = QPushButton("▶  " + self.translate_text("Lire"))
            pause_btn = QPushButton("⏸  " + self.translate_text("Pause"))
            stop_btn  = QPushButton("⏹  " + self.translate_text("Arrêter"))
            for btn in (play_btn, pause_btn, stop_btn):
                btn.setFixedHeight(34)
                hbox.addWidget(btn)

            vol_lbl = QLabel("🔊")
            vol_slider = QSlider(Qt.Horizontal)
            vol_slider.setRange(0, 100)
            vol_slider.setValue(80)
            vol_slider.setFixedWidth(90)
            hbox.addStretch()
            hbox.addWidget(vol_lbl)
            hbox.addWidget(vol_slider)
            vbox.addWidget(ctrl)

            play_btn.clicked.connect(player.play)
            pause_btn.clicked.connect(player.pause)
            stop_btn.clicked.connect(player.stop)
            vol_slider.valueChanged.connect(lambda v: audio_out.setVolume(v / 100))

            def _on_duration(dur):
                seek_slider.setRange(0, dur)
                secs = dur // 1000
                time_lbl.setText(f"0:00 / {secs//60}:{secs%60:02d}")

            def _on_position(pos):
                seek_slider.blockSignals(True)
                seek_slider.setValue(pos)
                seek_slider.blockSignals(False)
                dur = player.duration()
                def _fmt(ms):
                    s = ms // 1000
                    return f"{s//60}:{s%60:02d}"
                time_lbl.setText(f"{_fmt(pos)} / {_fmt(dur)}")

            def _on_seek(val):
                player.setPosition(val)

            player.durationChanged.connect(_on_duration)
            player.positionChanged.connect(_on_position)
            seek_slider.sliderMoved.connect(_on_seek)

            vbox.addStretch()
            self._clear_content()
            self.content_layout.addWidget(container)

        except ImportError as e:
            self._show_error(f"PySide6.QtMultimedia manquant: {e}")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur audio:')} {e}")

    def preview_video(self):
        try:
            from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
            from PySide6.QtMultimediaWidgets import QVideoWidget
            from PySide6.QtCore import QUrl
            from PySide6.QtWidgets import QSlider, QHBoxLayout, QVBoxLayout, QWidget, QLabel, QPushButton

            player = QMediaPlayer()
            audio_out = QAudioOutput()
            player.setAudioOutput(audio_out)
            audio_out.setVolume(0.8)
            self._media_player = player
            self._audio_output = audio_out

            video_widget = QVideoWidget()
            video_widget.setMinimumHeight(300)
            player.setVideoOutput(video_widget)
            player.setSource(QUrl.fromLocalFile(str(Path(self.file_path).resolve())))

            container = QWidget()
            vbox = QVBoxLayout(container)
            vbox.setSpacing(6)
            vbox.addWidget(video_widget, 1)

            time_lbl = QLabel("0:00 / 0:00")
            time_lbl.setAlignment(Qt.AlignCenter)
            time_lbl.setStyleSheet("font-size: 11px; color: gray;")
            seek_slider = QSlider(Qt.Horizontal)
            seek_slider.setRange(0, 0)
            vbox.addWidget(seek_slider)
            vbox.addWidget(time_lbl)

            ctrl = QWidget()
            hbox = QHBoxLayout(ctrl)
            hbox.setSpacing(10)

            play_btn  = QPushButton("▶  " + self.translate_text("Lire"))
            pause_btn = QPushButton("⏸  " + self.translate_text("Pause"))
            stop_btn  = QPushButton("⏹  " + self.translate_text("Arrêter"))
            for btn in (play_btn, pause_btn, stop_btn):
                btn.setFixedHeight(32)
                hbox.addWidget(btn)

            vol_lbl = QLabel("🔊")
            vol_slider = QSlider(Qt.Horizontal)
            vol_slider.setRange(0, 100)
            vol_slider.setValue(80)
            vol_slider.setFixedWidth(80)
            hbox.addStretch()
            hbox.addWidget(vol_lbl)
            hbox.addWidget(vol_slider)
            vbox.addWidget(ctrl)

            play_btn.clicked.connect(player.play)
            pause_btn.clicked.connect(player.pause)
            stop_btn.clicked.connect(player.stop)
            vol_slider.valueChanged.connect(lambda v: audio_out.setVolume(v / 100))

            def _on_duration(dur):
                seek_slider.setRange(0, dur)

            def _on_position(pos):
                seek_slider.blockSignals(True)
                seek_slider.setValue(pos)
                seek_slider.blockSignals(False)
                def _fmt(ms):
                    s = ms // 1000
                    return f"{s//60}:{s%60:02d}"
                time_lbl.setText(f"{_fmt(pos)} / {_fmt(player.duration())}")

            player.durationChanged.connect(_on_duration)
            player.positionChanged.connect(_on_position)
            seek_slider.sliderMoved.connect(player.setPosition)

            self._clear_content()
            self.content_layout.addWidget(container)
            self.resize(900, 620)

        except ImportError as e:
            self._show_error(f"PySide6.QtMultimediaWidgets manquant: {e}")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur vidéo:')} {e}")

    def preview_unsupported(self):
        ext = Path(self.file_path).suffix.upper() or "?"
        self.preview_label.setParent(self.content_area)
        self.content_layout.addWidget(self.preview_label)
        self.preview_label.setText(
            f"🚫  {self.translate_text('Aperçu non disponible pour ce type de fichier')}\n({ext})"
        )

    def closeEvent(self, event):
        if self._media_player:
            try:
                self._media_player.stop()
            except Exception:
                pass
        super().closeEvent(event)
