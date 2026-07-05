"""
context_menu/window.py
Quick Convert popup for Windows Shell Integration for File Converter Pro

Launched via:
    FileConverterPro.exe --context-menu --files "path1" "path2" ...  --conversion-type [type]
"""

import os
import sys
import time
import tempfile
from pathlib import Path
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QComboBox, QProgressBar, QFrame, QApplication
)
from PySide6.QtGui  import QIcon
from PySide6.QtCore import Qt, QThread, Signal, QTimer

from qss_helpers import _load_qss
from context_menu.formats import CONVERSION_MAP, LABELS, MERGE_DISPATCH

_COLLECT_MIN_SEC    = 2.0
_COLLECT_STABLE_SEC = 0.5
_COLLECT_MAX_SEC    = 8.0
_COLLECT_POLL_SEC   = 0.05
_COLLECT_EPOCH_SEC  = 15.0

def _collect_staged_files(conversion_type: str, single_file: str) -> list[str]:
    """
    Windows spawns one process per selected file for context menu commands.
    """
    epoch   = int(time.time() / _COLLECT_EPOCH_SEC)
    key     = f"fcp_{conversion_type}_{epoch}"
    staging = os.path.join(tempfile.gettempdir(), f"{key}.txt")
    lock    = os.path.join(tempfile.gettempdir(), f"{key}.lock")

    with open(staging, "a", encoding="utf-8") as fh:
        fh.write(single_file + "\n")

    time.sleep(_COLLECT_MIN_SEC)

    deadline     = time.time() + (_COLLECT_MAX_SEC - _COLLECT_MIN_SEC)
    last_size    = -1
    stable_since = time.time()

    while time.time() < deadline:
        time.sleep(_COLLECT_POLL_SEC)
        try:
            size = os.path.getsize(staging)
        except OSError:
            size = last_size
        if size != last_size:
            last_size    = size
            stable_since = time.time()
        elif time.time() - stable_since >= _COLLECT_STABLE_SEC:
            break

    try:
        fd = os.open(lock, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
        os.close(fd)
    except FileExistsError:
        return []

    time.sleep(0.2)

    try:
        with open(staging, "r", encoding="utf-8") as fh:
            return [l.strip() for l in fh if l.strip() and os.path.isfile(l.strip())]
    except Exception:
        return [single_file]

AUTO_CLOSE_DELAY_MS = 1800


def _sort_files(files: list, order: str) -> list:
    if order == "none" or not order:
        return files
    if order == "alphabetical_asc":
        return sorted(files, key=lambda f: Path(f).name.lower())
    if order == "alphabetical_desc":
        return sorted(files, key=lambda f: Path(f).name.lower(), reverse=True)
    if order == "numerical_asc":
        import re
        return sorted(files, key=lambda f: [
            int(t) if t.isdigit() else t.lower()
            for t in re.split(r"(\d+)", Path(f).name)
        ])
    if order == "numerical_desc":
        import re
        return sorted(files, key=lambda f: [
            int(t) if t.isdigit() else t.lower()
            for t in re.split(r"(\d+)", Path(f).name)
        ], reverse=True)
    if order == "date_asc":
        return sorted(files, key=lambda f: Path(f).stat().st_mtime)
    if order == "date_desc":
        return sorted(files, key=lambda f: Path(f).stat().st_mtime, reverse=True)
    return files


def _convert_image_to_pdf(src: str, dst_dir: str) -> None:
    import fitz
    out = os.path.join(dst_dir, f"{Path(src).stem}.pdf")
    pdf = fitz.open()
    img = fitz.open(src)
    rect = img[0].rect
    page = pdf.new_page(width=rect.width, height=rect.height)
    page.insert_image(rect, filename=src)
    img.close()
    pdf.save(out)
    pdf.close()


def _merge_images_to_pdf(files: list, dst_dir: str) -> None:
    import fitz
    pdf = fitz.open()
    for src in files:
        img = fitz.open(src)
        rect = img[0].rect
        page = pdf.new_page(width=rect.width, height=rect.height)
        page.insert_image(rect, filename=src)
        img.close()
    import time as _time
    out = os.path.join(dst_dir, f"merged_images_{_time.strftime('%Y%m%d_%H%M%S')}.pdf")
    pdf.save(out)
    pdf.close()


def _merge_pdfs(files: list, dst_dir: str, order: str) -> None:
    import fitz
    merged = fitz.open()
    import time as _time
    for f in _sort_files(files, order):
        doc = fitz.open(f)
        merged.insert_pdf(doc)
        doc.close()
    merged.save(os.path.join(dst_dir, f"merged_{_time.strftime('%Y%m%d_%H%M%S')}.pdf"))
    merged.close()


def _merge_docx(files: list, dst_dir: str, order: str) -> None:
    from docxcompose.composer import Composer
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import time as _time

    sorted_files = _sort_files(files, order)
    merged = Document(sorted_files[0])
    composer = Composer(merged)

    for f in sorted_files[1:]:
        doc = Document(f)
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r.append(br)
        p.append(r)
        doc.element.body.insert(0, p)
        composer.append(doc)

    merged.save(os.path.join(dst_dir, f"merged_{_time.strftime('%Y%m%d_%H%M%S')}.docx"))


def _convert_docx_to_pdf(src: str, dst_dir: str) -> None:
    out     = os.path.join(dst_dir, f"{Path(src).stem}.pdf")
    src_abs = os.path.abspath(src)
    out_abs = os.path.abspath(out)

    try:
        from docx2pdf import convert
        convert(src_abs, out_abs); return
    except Exception as e:
        print(f"[docx→pdf] docx2pdf failed: {e}")

    try:
        import comtypes.client
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(src_abs)
        doc.SaveAs(out_abs, FileFormat=17)
        doc.Close(False); word.Quit(); return
    except Exception as e:
        print(f"[docx→pdf] COM failed: {e}")

    try:
        import subprocess, shutil
        candidates = [
            "soffice", "libreoffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        exe = next((c for c in candidates if shutil.which(c) or os.path.exists(c)), None)
        if exe:
            subprocess.run(
                [exe, "--headless", "--convert-to", "pdf", "--outdir", dst_dir, src_abs],
                check=True, capture_output=True, timeout=60
            ); return
    except Exception as e:
        print(f"[docx→pdf] LibreOffice failed: {e}")

    try:
        from docx import Document as DocxDoc
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        doc = DocxDoc(src_abs); styles = getSampleStyleSheet(); story = []
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles["Normal"]))
                story.append(Spacer(1, 6))
        SimpleDocTemplate(out_abs, pagesize=A4).build(story); return
    except Exception as e:
        print(f"[docx→pdf] reportlab failed: {e}")

    raise RuntimeError("All DOCX→PDF methods failed.")


def _convert_pdf_to_docx(src: str, dst_dir: str) -> None:
    import sys
    import threading

    out     = os.path.join(dst_dir, f"{Path(src).stem}.docx")
    src_abs = os.path.abspath(src)
    out_abs = os.path.abspath(out)

    # Tier 1: Word COM (best quality, Windows only)
    if sys.platform == "win32":
        try:
            import winreg
            winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, "Word.Application")
            _word_available = True
        except (OSError, ImportError):
            _word_available = False
            print("[pdf→docx] Word not found in registry — skipping COM tier")

        if _word_available:
            result     = {"ok": False}
            stop_event = threading.Event()

            def _dialog_dismisser():
                import ctypes, ctypes.wintypes, time as _t
                user32   = ctypes.windll.user32
                BM_CLICK = 0x00F5

                DialogProc = ctypes.WINFUNCTYPE(
                    ctypes.c_bool,
                    ctypes.wintypes.HWND,
                    ctypes.wintypes.LPARAM,
                )

                def _find_ok_button(hwnd_dialog):
                    found = ctypes.wintypes.HWND(0)

                    @DialogProc
                    def _enum(hwnd_child, _lp):
                        buf = ctypes.create_unicode_buffer(64)
                        user32.GetWindowTextW(hwnd_child, buf, 64)
                        if buf.value.strip().upper() in ("OK", "O&K"):
                            found.value = hwnd_child
                            return False
                        if user32.GetDlgCtrlID(hwnd_child) == 1:
                            found.value = hwnd_child
                            return False
                        return True

                    user32.EnumChildWindows(hwnd_dialog, _enum, 0)
                    return found.value or None

                dialog_classes = ["#32770", "bosa_sdm_msword"]
                target_title   = "Microsoft Word"

                while not stop_event.is_set():
                    for cls in dialog_classes:
                        hwnd = user32.FindWindowW(cls, None)
                        if hwnd:
                            ok_btn = _find_ok_button(hwnd)
                            if ok_btn:
                                user32.SendMessageW(ok_btn, BM_CLICK, 0, 0)
                            else:
                                user32.SendMessageW(hwnd, 0x0111, 1, 0)
                    hwnd = user32.FindWindowW(None, target_title)
                    if hwnd:
                        ok_btn = _find_ok_button(hwnd)
                        if ok_btn:
                            user32.SendMessageW(ok_btn, BM_CLICK, 0, 0)
                        else:
                            user32.SendMessageW(hwnd, 0x0111, 1, 0)
                    _t.sleep(0.3)

            def _com_worker():
                word = None
                doc  = None
                try:
                    import pythoncom, comtypes.client
                    pythoncom.CoInitialize()
                    word = comtypes.client.CreateObject("Word.Application")
                    word.Visible           = False
                    word.DisplayAlerts     = 0
                    word.AutomationSecurity = 3
                    doc = word.Documents.Open(
                        src_abs,
                        ConfirmConversions = False,
                        ReadOnly           = True,
                        AddToRecentFiles   = False,
                        NoEncodingDialog   = True,
                    )
                    doc.SaveAs2(out_abs, FileFormat=16)
                    result["ok"] = True
                    print(f"[pdf→docx] Word COM success: {Path(src_abs).name}")
                except Exception as e:
                    print(f"[pdf→docx] COM failed: {e}")
                finally:
                    if doc is not None:
                        try: doc.Close(False)
                        except Exception: pass
                    if word is not None:
                        try: word.Quit()
                        except Exception: pass
                    try:
                        import pythoncom as _pc; _pc.CoUninitialize()
                    except Exception: pass

            dismisser = threading.Thread(target=_dialog_dismisser, daemon=True)
            dismisser.start()
            worker = threading.Thread(target=_com_worker, daemon=True)
            worker.start()
            worker.join(timeout=60)
            stop_event.set()

            if result["ok"]:
                return
            print("[pdf→docx] COM timed out or failed — falling back")

    # Tier 2: pdf2docx
    try:
        from pdf2docx import Converter
        cv = Converter(src_abs); cv.convert(out_abs, parse_drawing=True); cv.close(); return
    except Exception as e:
        print(f"[pdf→docx] pdf2docx failed: {e}")

    # Tier 3: fitz + python-docx (text only)
    try:
        import fitz
        from docx import Document as DocxDoc
        pdf_doc  = fitz.open(src_abs)
        word_doc = DocxDoc()
        for page_num in range(len(pdf_doc)):
            text = pdf_doc.load_page(page_num).get_text("text")
            for line in text.split("\n"):
                if line.strip():
                    word_doc.add_paragraph(line.strip())
            if page_num < len(pdf_doc) - 1:
                word_doc.add_page_break()
        pdf_doc.close(); word_doc.save(out_abs); return
    except Exception as e:
        print(f"[pdf→docx] fitz failed: {e}")

    raise RuntimeError("All PDF→DOCX methods failed.")


class ConversionThread(QThread):
    done = Signal(bool, str, str)

    def __init__(self, files: list, conversion_type: str):
        super().__init__()
        self.files           = files
        self.conversion_type = conversion_type

    def run(self):
        try:
            base = os.path.dirname(
                sys.executable if getattr(sys, "frozen", False)
                else os.path.abspath(os.path.join(__file__, "..", ".."))
            )
            if base not in sys.path:
                sys.path.insert(0, base)


            dst_dir = os.path.dirname(self.files[0])

            if self.conversion_type == "images_to_pdf_merged":
                _merge_images_to_pdf(self.files, dst_dir)
                self.done.emit(True, "Images merged into PDF", dst_dir)
                return

            if self.conversion_type in MERGE_DISPATCH:
                kind, order = MERGE_DISPATCH[self.conversion_type]
                if kind == "pdf":
                    _merge_pdfs(self.files, dst_dir, order)
                else:
                    _merge_docx(self.files, dst_dir, order)
                self.done.emit(True, "Files merged successfully", dst_dir)
                return

            errors = []
            out_dir = ""
            for f in self.files:
                f_dst_dir = os.path.dirname(f)
                try:
                    if self.conversion_type == "image_to_pdf":
                        _convert_image_to_pdf(f, f_dst_dir)
                    elif self.conversion_type == "docx_to_pdf":
                        _convert_docx_to_pdf(f, f_dst_dir)
                    elif self.conversion_type == "pdf_to_docx":
                        _convert_pdf_to_docx(f, f_dst_dir)
                    else:
                        from converter.converters import AdvancedConverterEngine
                        result = AdvancedConverterEngine().convert(
                            self.conversion_type, f, f_dst_dir
                        )
                        if not result.success:
                            raise Exception(result.error)
                    out_dir = f_dst_dir
                except Exception as e:
                    errors.append(f"{os.path.basename(f)}: {e}")

            if errors:
                self.done.emit(False, "\n".join(errors), "")
            else:
                n = len(self.files)
                self.done.emit(True, f"{n} file{'s' if n > 1 else ''} converted", out_dir)

        except Exception as e:
            self.done.emit(False, str(e), "")


class QuickConvertWindow(QWidget):
    """Floating window launched from the Windows context menu."""

    def __init__(self, files: list, conversion_type: str | None = None):
        super().__init__()
        self.files = [f for f in files if os.path.isfile(f)]
        self._converting = False

        self.ext = (
            os.path.splitext(self.files[0])[1].lstrip(".").lower()
            if self.files else ""
        )
        self.forced_conversion = conversion_type
        self.conversion_types  = (
            [conversion_type] if conversion_type
            else CONVERSION_MAP.get(self.ext, [])
        )

        self.setWindowTitle("File Converter Pro")
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint | Qt.Tool)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setMinimumWidth(310)
        self.setMaximumWidth(310)

        self._set_icon()
        self._build_ui()
        self._center_on_screen()

    def _set_icon(self):
        from utils import get_icon_path
        icon_path = get_icon_path("icon.ico")
        if icon_path and os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        self.card = QFrame()
        self.card.setObjectName("card")
        self.card.setStyleSheet(_load_qss("window_card.qss"))

        inner = QVBoxLayout(self.card)
        inner.setSpacing(10)
        inner.setContentsMargins(18, 16, 18, 16)
        self._build_ui_inner(inner)

        root.addWidget(self.card)
        self.progress.show()
        self.adjustSize()
        self.setFixedHeight(self.sizeHint().height())
        self.progress.hide()

    def _build_ui_inner(self, inner: QVBoxLayout):
        title = QLabel("Quick Convert")
        title.setObjectName("title")
        inner.addWidget(title)

        n = len(self.files)
        if n == 1:
            file_text = os.path.basename(self.files[0])
        elif n <= 4:
            file_text = "\n".join(f"• {os.path.basename(f)}" for f in self.files)
        else:
            previews  = "\n".join(f"• {os.path.basename(f)}" for f in self.files[:3])
            file_text = f"{previews}\n• … and {n - 3} more"

        file_label = QLabel(file_text)
        file_label.setObjectName("fileList")
        file_label.setWordWrap(True)
        inner.addWidget(file_label)

        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet("background-color: #313244; max-height: 1px;")
        inner.addWidget(sep)

        self.combo = None

        if not self.conversion_types:
            msg = QLabel(f"No conversion available for .{self.ext or '?'}")
            msg.setObjectName("fileList")
            msg.setWordWrap(True)
            inner.addWidget(msg)
        elif self.forced_conversion:
            label = QLabel(LABELS.get(self.forced_conversion, self.forced_conversion))
            label.setObjectName("conversionLabel")
            inner.addWidget(label)
        else:
            self.combo = QComboBox()
            for ct in self.conversion_types:
                self.combo.addItem(LABELS.get(ct, ct), ct)
            inner.addWidget(self.combo)

        self.progress = QProgressBar()
        self.progress.setRange(0, 0)
        self.progress.hide()
        inner.addWidget(self.progress)

        self.status_label = QLabel("")
        self.status_label.setObjectName("statusOk")
        self.status_label.setWordWrap(True)
        self.status_label.hide()
        inner.addWidget(self.status_label)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.setObjectName("cancelBtn")
        self.cancel_btn.clicked.connect(self.close)

        self.convert_btn = QPushButton("Convert")
        self.convert_btn.setObjectName("convertBtn")
        self.convert_btn.setEnabled(bool(self.conversion_types))
        self.convert_btn.clicked.connect(self._start_conversion)

        btn_row.addWidget(self.cancel_btn)
        btn_row.addWidget(self.convert_btn)
        inner.addLayout(btn_row)

    def _center_on_screen(self):
        geo = QApplication.primaryScreen().availableGeometry()
        self.move(
            geo.center().x() - self.width()  // 2,
            geo.center().y() - self.height() // 2,
        )

    def _start_conversion(self):
        conversion_type = (
            self.forced_conversion
            if self.forced_conversion
            else self.combo.currentData()
        )
        self._converting = True
        self.convert_btn.setEnabled(False)
        self.cancel_btn.setEnabled(False)
        self.progress.show()
        self.status_label.hide()

        self._worker_thread = ConversionThread(self.files, conversion_type)
        self._worker_thread.done.connect(self._on_done)
        self._worker_thread.start()

    def _on_done(self, success: bool, message: str, out_dir: str):
        self._converting = False
        self.progress.hide()
        self.status_label.show()
        self.adjustSize()
        self.setFixedHeight(self.sizeHint().height())

        if success:
            self.status_label.setObjectName("statusOk")
            self.status_label.setText(f"✓ {message}")
            QTimer.singleShot(AUTO_CLOSE_DELAY_MS, QApplication.quit)
        else:
            self.status_label.setObjectName("statusErr")
            self.status_label.setText(f"✗ {message}")
            self.convert_btn.setText("Retry")
            self.convert_btn.setEnabled(True)
            self.cancel_btn.setEnabled(True)
            self.convert_btn.clicked.disconnect()
            self.convert_btn.clicked.connect(self._start_conversion)

        self.status_label.style().unpolish(self.status_label)
        self.status_label.style().polish(self.status_label)

    def mousePressEvent(self, e):
        if e.button() == Qt.LeftButton:
            self._drag_pos = e.globalPosition().toPoint() - self.frameGeometry().topLeft()

    def mouseMoveEvent(self, e):
        if e.buttons() == Qt.LeftButton and hasattr(self, "_drag_pos"):
            self.move(e.globalPosition().toPoint() - self._drag_pos)

    def closeEvent(self, e):
        e.accept()
        if hasattr(self, "_worker_thread") and self._worker_thread.isRunning():
            self._worker_thread.quit()
            self._worker_thread.wait(2000)
            os._exit(0)
        else:
            QApplication.quit()


def run_context_menu(files: list, conversion_type: str | None = None) -> None:
    """Called from main.py when --context-menu flag is detected."""

    if conversion_type and len(files) == 1:
        collected = _collect_staged_files(conversion_type, files[0])
        if not collected:
            sys.exit(0)
        files = collected

    app = QApplication.instance() or QApplication(sys.argv)
    win = QuickConvertWindow(files, conversion_type=conversion_type)
    win.show()

    if conversion_type:
        QTimer.singleShot(100, win._start_conversion)

    sys.exit(app.exec())